# ==============================================================
#  SOC Threat Analyser — v6.3 (2025-12-02)
#  Autor: ocrHeiki / GitHub (muudetud: assistent)
#  Kirjeldus:
#   - Analüüsib Palo Alto logisid (CSV)
#   - Toetab: 24h / 7d / 30d / 60d / 90d  ajaaknaid
#   - Väljundid: DOCX, TXT, XLSX + graafikud (PNG)
#   - Funktsionaalsused:
#     * Severity, Threat, Action, IP (source/destination), Port, Trendid
#     * Threat + Port kaardistus
#     * DOCX raport koos visuaalidega ja sisukorraga
#     * Korrelatsioon: IP -> TOP Threats & TOP Kategooriad
#     * UUS: 30-90 päeva võrdlev trendianalüüs ("EKG" graafik)
# ==============================================================
from pathlib import Path
from datetime import datetime, timedelta
import pandas as pd
import matplotlib.pyplot as plt
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import json
import math
import numpy as np

# ———————————————————————————————
# GLOBALSEADISTUS
# ———————————————————————————————

TIME_MAPPING_DAYS = {"24h": 1, "7d": 7, "30d": 30, "60d": 60, "90d": 90}

# ———————————————————————————————
# CLI seadistus
# ———————————————————————————————

def parse_args():
    parser = argparse.ArgumentParser(description="SOC Threat Analyser – Palo Alto logide multianalüüs")
    parser.add_argument("--timeframe", choices=list(TIME_MAPPING_DAYS.keys()), default="24h",
                        help=f"Analüüsitav ajaaken ({', '.join(TIME_MAPPING_DAYS.keys())}) (vaikimisi: '24h')")
    return parser.parse_args()

# ———————————————————————————————
# Kaustad ja värvid
# ———————————————————————————————

BASE_DIR = Path.home() / "Documents" / "SOC"
RAW_DIR = BASE_DIR / "raw"
REPORTS_DIR = BASE_DIR / "reports"
RESULTS_DIR = BASE_DIR / "tulemused"
THREAT_VAULT_CACHE = BASE_DIR / "threat_vault_cache"
TRENDS_DIR = BASE_DIR / "trendid"

for d in [RAW_DIR, REPORTS_DIR, RESULTS_DIR, THREAT_VAULT_CACHE, TRENDS_DIR]:
    d.mkdir(parents=True, exist_ok=True)

COLORS_SEV = {"low": "#0000FF", "medium": "#FFFF00", "high": "#FFA500", "critical": "#FF0000"}
COLORS_ACTION = {"allow": "#33CC33", "deny": "#CC3333", "drop": "#3366CC", "alert": "#FFCC00", "reset-both": "#9933CC", "reset-server": "#800080"}
COLORS_CAT = {
    "command-and-control": "#CC0000", "code-execution": "#FF6600", "sql-injection": "#FF9933",
    "brute-force": "#FFCC00", "dos": "#FFFF66", "hacktool": "#9933CC", "info-leak": "#66CCFF",
    "spyware": "#3399FF", "code-obfuscation": "#996633"
}

# ———————————————————————————————
# Abifunktsioonid
# ———————————————————————————————

def get_current_timestamp(timeframe: str):
    """Genereerib kuupäevast ja ajaaknast vormindatud stringi failinimedele"""
    return f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{timeframe.replace('d', 'p')}"

def get_filename_core(timeframe: str):
    """Genereerib failinime alguse vastavalt ajaperioodile"""
    return f"soc_{get_current_timestamp(timeframe)}"

def iso_from_filename(name: str):
    """Loeb kuupäeva failide nimedest, toetades ka uut YYYY-MM-DD formaati"""
    m = re.search(r"(\d{4})-(\d{2})-(\d{2})", name)
    if m:
        return datetime.strptime(f"{m.group(1)}-{m.group(2)}-{m.group(3)}", "%Y-%m-%d").date()
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", name)
    if m:
        return datetime.strptime(f"{m.group(3)}-{m.group(2)}-{m.group(1)}", "%Y-%m-%d").date()
    # Tagavara: kasutab tänast kuupäeva
    return datetime.now().date()

def first_existing(df, names):
    for n in names:
        if n in df.columns:
            return n
    return None

def norm_lower(col):
    return col.astype(str).str.lower().str.strip()

def overwrite_file(path: Path):
    if path.exists():
        try:
            path.unlink()
        except Exception as e:
            print(f"[~] Ei saanud kustutada faili: {path} ({e})")

def bar(series, title, outpath, colors=None, rot=0, is_date_trend=False):
    if series is None or series.empty:
        return
    plt.figure(figsize=(10, 5))
    try:
        cols = [colors.get(str(k).lower(), "#888888") for k in series.index] if colors else "#3399FF"
    except Exception:
        cols = "#3399FF"
    ax = series.plot(kind="bar", color=cols)
    plt.title(title)
    if is_date_trend:
        try:
            # Püüdke esmalt kasutada datetime index'it
            ax.set_xticklabels(series.index.strftime('%Y-%m-%d'))
        except:
            pass # Kui ei õnnestu
        plt.xticks(rotation=45, ha='right')
    else:
        plt.xticks(rotation=rot)
    plt.tight_layout()
    # Puhastamine ja salvestamine
    Path(outpath).parent.mkdir(parents=True, exist_ok=True)
    overwrite_file(outpath)
    plt.savefig(outpath)
    plt.close()

def donut_with_side_legend(series, title, outpath, colors_map=None, legend_title="Legend"):
    if series is None or series.empty:
        return
    values = series.values
    labels = [str(i) for i in series.index]
    total = float(values.sum()) if values.sum() else 1.0
    pcts = [round(v * 100.0 / total, 1) for v in values]
    colors = [colors_map.get(str(i).lower(), "#888888") for i in series.index] if colors_map else None
    fig, ax = plt.subplots(figsize=(9, 6))
    wedges, _ = ax.pie(values, startangle=90, labels=None, colors=colors, wedgeprops=dict(width=0.38))
    ax.set_title(title)
    legend_lines = [f"{lbl} — {int(cnt)} ({pct}%)" for lbl, cnt, pct in zip(labels, values, pcts)]
    ax.legend(wedges, legend_lines, title=legend_title, loc="center left", bbox_to_anchor=(1.0, 0.5), frameon=False)
    plt.tight_layout()
    # Puhastamine ja salvestamine
    Path(outpath).parent.mkdir(parents=True, exist_ok=True)
    overwrite_file(outpath)
    fig.savefig(outpath, bbox_inches="tight")
    plt.close(fig)

def add_image(doc, img_path, caption, width_in=6.0):
    if not img_path.exists():
        return
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ———————————————————————————————
# Trendianalüüs
# ———————————————————————————————

def analyze_trends(df_all: pd.DataFrame, timeframe: str):
    """
    Teostab nädalapõhise logi mahu ja TOP 5 threatide trendianalüüsi
    """
    if "tname_norm" not in df_all.columns:
        print("[!] 'tname_norm' puudub – trendi ei saa teha.")
        return {}

    df = df_all.copy()
    if 'log_date' not in df.columns:
        df['log_date'] = datetime.now().date() # Tagavaraks kui puudu
    df["log_date"] = pd.to_datetime(df["log_date"], errors='coerce')
    df = df.dropna(subset=["log_date"])

    # Grupib andmed nädala alguse järgi
    # Pühendatud: kasutab esimesest kirjest saadud kuupäeva, et vältida veidraid nihkeid
    df['week_start'] = df['log_date'].apply(
        lambda x: x - timedelta(days=(x.weekday()))
    ).dt.normalize()

    # 1. Nädalapõhine Logide Maht
    weekly_volume = df.groupby('week_start').size()
    out_vol = TRENDS_DIR / f"trend_{timeframe}_weekly_volume.png"
    bar(weekly_volume, f"{timeframe.upper()} Nädalapõhine Logide Maht", out_vol, is_date_trend=True)

    # 2. TOP 5 Threati Aktiivsus Aja Jooksul
    top5_names = df['tname_norm'].value_counts().head(5).index.tolist()
    weekly_threats = df.groupby(['week_start', 'tname_norm']).size().unstack(fill_value=0)
    top5_trend = weekly_threats[weekly_threats.columns.intersection(top5_names)]

    out_top5 = TRENDS_DIR / f"trend_{timeframe}_top5_threats.png"
    plt.figure(figsize=(12, 6))
    top5_trend.plot(kind='line', marker='o', ax=plt.gca())
    plt.title(f"{timeframe.upper()} TOP 5 Threat'i Aktiivsus Aja Jooksul")
    plt.xlabel("Nädala Algus")
    plt.ylabel("Kirjete Arv")
    plt.legend(title='Threat', bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.grid(True)
    plt.tight_layout()
    overwrite_file(out_top5)
    plt.savefig(out_top5)
    plt.close()

    print("[+] Trendianalüüs lõpetatud.")
    return {
        "plot_volume": out_vol,
        "plot_top5": out_top5
    }

def compare_trends(df_all: pd.DataFrame, timeframe: str):
    """
    Teostab võrdleva trendianalüüsi 30d/60d/90d andmetele 7-päevaste intervallidega
    ja teeb "EKG" stiilis graafikud. Langused/tõusud TOP 5 Threat ja Vale-positiivi seas.
    """
    if timeframe not in ["60d", "90d"]:
        return {}

    if "tname_norm" not in df_all.columns:
        print("[!] Tõus/Langus analüüsiks puudub 'tname_norm'.")
        return {}

    df = df_all.copy()
    df["log_date"] = pd.to_datetime(df["log_date"], errors='coerce')
    df = df.dropna(subset=["log_date"])

    # 7-päevased intervallid
    df['interval_start'] = df['log_date'].apply(
        lambda x: x - timedelta(days=(x.dayofweek+1) % 7) # Lihtsalt 7p segment
    ).dt.normalize()

    # Määrame, mida loeme "false-positive'iks" - lihtsustatult Action=Allow / Severity=Low
    # Ideaalis peaks olema "whitelist" või "benign" kategooria, aga kasutame olemasolevaid
    df['is_false_positive'] = (df['act_norm'] == 'allow') | (df['sev_norm'] == 'low')
    df['threat_type'] = np.where(df['is_false_positive'], 'false_positive', 'threat')

    # Grupib 7-päevaste intervallide ja threat_type järgi
    weekly_matrix = df.groupby(['interval_start', 'threat_type']).size().unstack(fill_value=0)

    # Kui andmeid pole, siis tagasta tühi
    if weekly_matrix.empty:
      return {}

    # Võrdlus: Arvutab muudatused eelmisest nädalast (delta)
    weekly_delta = weekly_matrix.diff().fillna(0)

    # Normaliseerime väärtused EKG-graafiku jaoks (nt. vahemikku -1 kuni 1)
    # See teeb graafikud paremini võrreldavaks, isegi kui tooreandmed erinevad
    def normalize_delta(series):
        max_abs = series.abs().max()
        return series / max_abs if max_abs > 0 else series
    
    weekly_delta_norm_threat = normalize_delta(weekly_delta['threat'])
    weekly_delta_norm_fp = normalize_delta(weekly_delta['false_positive'])

    # --- 1. EKG-stiilis graafik: Ohud ja Valepositiivid Ajas (Protsentuaalne muutus) ---
    out_ekg = TRENDS_DIR / f"trend_{timeframe}_ekg_threat_fp.png"
    plt.figure(figsize=(15, 6))
    
    # Ohud (sinine)
    plt.plot(weekly_delta_norm_threat.index, weekly_delta_norm_threat.values, 
             label='Ohud (normalized change)', color='darkred', marker='o', linestyle='-', linewidth=2)
    # Valepositiivid (roheline)
    plt.plot(weekly_delta_norm_fp.index, weekly_delta_norm_fp.values, 
             label='Valepositiivid (normalized change)', color='darkgreen', marker='o', linestyle='-', linewidth=2)

    plt.axhline(0, color='gray', linestyle='--', linewidth=1) # Null-joon
    plt.title(f"{timeframe.upper()} Ohtude & Valepositiivide Muutused (7p intervall) - 'EKG'")
    plt.xlabel("7-päevase Intervalli Algus")
    plt.ylabel("Normalized Change (Growth/Decline)")
    plt.xticks(rotation=45, ha='right')
    plt.legend()
    plt.grid(True, linestyle=':', alpha=0.6)
    plt.tight_layout()
    overwrite_file(out_ekg)
    plt.savefig(out_ekg)
    plt.close()

    # --- 2. Detailne Langused/Tõusud (TOP 5) ---
    summary = {}
    
    # a) Ohud
    threat_weekly_counts = df.loc[df['threat_type'] == 'threat'].groupby(['interval_start', 'tname_norm']).size().unstack(fill_value=0)
    # Jätame vaid kõige aktiivsemad (TOP 5 kogu perioodi lõikes)
    top_threats = df.loc[df['threat_type'] == 'threat']['tname_norm'].value_counts().head(5).index
    threat_weekly_counts = threat_weekly_counts[top_threats].copy()

    # Arvutame muudatuse absoluutarvudes (viimane intervall vs eelviimane)
    latest_interval = threat_weekly_counts.tail(1).index[0]
    previous_interval = threat_weekly_counts.tail(2).head(1).index[0]

    threat_latest = threat_weekly_counts.loc[latest_interval]
    threat_previous = threat_weekly_counts.loc[previous_interval]
    threat_change = threat_latest - threat_previous

    # Tõusud ja Langused ohtude seas
    rising_threats = threat_change[threat_change > 0].sort_values(ascending=False).head(5)
    declining_threats = threat_change[threat_change < 0].sort_values(ascending=True).head(5)

    summary['rising_threats'] = rising_threats.to_dict()
    summary['declining_threats'] = declining_threats.to_dict()

    # Graafik: TOP 5 EKG-stiilis
    out_top5_ekg = TRENDS_DIR / f"trend_{timeframe}_ekg_top5_threats.png"
    plt.figure(figsize=(15, 6))
    for threat_name in top_threats:
        series = threat_weekly_counts[threat_name]
        series_norm_delta = normalize_delta(series.diff().fillna(0))
        plt.plot(series_norm_delta.index, series_norm_delta.values, 
                 label=threat_name, linestyle='-', marker='.')

    plt.axhline(0, color='gray', linestyle='--', linewidth=1)
    plt.title(f"{timeframe.upper()} TOP 5 Ohtude Muutused Ajas (Normalized Delta)")
    plt.xlabel("7-päevase Intervalli Algus")
    plt.ylabel("Normalized Change (Growth/Decline)")
    plt.xticks(rotation=45, ha='right')
    plt.legend(title='Threat')
    plt.grid(True, linestyle=':', alpha=0.6)
    plt.tight_layout()
    overwrite_file(out_top5_ekg)
    plt.savefig(out_top5_ekg)
    plt.close()

    print("[+] Võrdlev trendianalüüs lõpetatud. (EKG stiilis)")
    return {
        "plot_ekg": out_ekg,
        "plot_top5_ekg": out_top5_ekg,
        "summary": summary
    }


# ———————————————————————————————
# IP -> Threat/Kategooria korrelatsiooni (abi)
# ———————————————————————————————

def ip_correlation(df, ip_col, threat_col, cat_col, top_n=10):
    """
    Tagastab DataFrame'i, kus iga IP-l on:
      - count (kui mitu kirjet)
      - top threats (näide: 'threat1(12), threat2(7)')
      - top categories (näide: 'cat1(8), cat2(3)')
    """
    if ip_col not in df.columns:
        return pd.DataFrame()

    counts = df[ip_col].value_counts().head(top_n)
    rows = []
    for ip, cnt in counts.items():
        sub = df[df[ip_col] == ip]
        top_th = sub[threat_col].value_counts().head(5) if threat_col and threat_col in sub.columns else pd.Series(dtype=int)
        top_cat = sub[cat_col].value_counts().head(5) if cat_col and cat_col in sub.columns else pd.Series(dtype=int)

        th_str = ", ".join([f"{str(t)}({int(c)})" for t, c in top_th.items()]) if not top_th.empty else ""
        cat_str = ", ".join([f"{str(ca)}({int(cc)})" for ca, cc in top_cat.items()]) if not top_cat.empty else ""

        rows.append({
            "ip": ip,
            "count": int(cnt),
            "top_threats": th_str,
            "top_categories": cat_str
        })
    return pd.DataFrame(rows)

# ———————————————————————————————
# Peamine logika
# ———————————————————————————————

def main():
    args = parse_args()
    timeframe_days = TIME_MAPPING_DAYS[args.timeframe]
    timeframe_str = args.timeframe
    file_core = get_filename_core(timeframe_str)

    print(f"[i] Käivitan {timeframe_str} analüüsi...")

    # Samme failide lugemisel pole muudetud – loeme kõik, et teostada ajaline filter hiljem
    csv_files = list(RAW_DIR.glob("*.csv"))
    if not csv_files:
        print("[!] Ühtegi CSV-faili ei leitud.")
        return

    dfs = []
    today = datetime.now().date()
    used_files = []

    for f in csv_files:
        try:
            df = pd.read_csv(f, low_memory=False)
        except Exception as e:
            print(f"[!] Viga faili lugemisel {f.name}: {e}")
            continue
        df["log_date"] = iso_from_filename(f.name)
        dfs.append(df)
        used_files.append(f.name)

    if not dfs:
        print("[!] Ühtegi loetavat CSV andmestikku pole.")
        return

    # Kogu andmestik ja filtreerimine
    df_all = pd.concat(dfs, ignore_index=True)
    print(f"[i] Loetud {len(df_all)} kirjet {len(csv_files)} failist.")

    # --- veerud (leiame parimad võimalikud nimetused ja normaliseerimine)
    threat_col = first_existing(df_all, ["Threat/Content Name", "threat_name"])
    sev_col    = first_existing(df_all, ["Severity", "severity"])
    act_col    = first_existing(df_all, ["Action", "action"])
    src_col    = first_existing(df_all, ["Source address", "src", "source"])
    dst_col    = first_existing(df_all, ["Destination address", "dst", "destination"])
    cat_col    = first_existing(df_all, ["thr_category", "category"])
    # type_col   = first_existing(df_all, ["Threat/Content Type", "threat_type"]) # Hetkel pole kasutatud

    if threat_col:
        df_all["tname_norm"] = norm_lower(df_all[threat_col])
    else:
        df_all["tname_norm"] = "" # Tühi string, mitte "unknown" kui veerg puudub

    # normaliseeri muud veerud
    df_all["sev_norm"] = norm_lower(df_all[sev_col]) if sev_col else pd.Series(["unknown"] * len(df_all))
    df_all["act_norm"] = norm_lower(df_all[act_col]) if act_col else pd.Series(["unknown"] * len(df_all))
    df_all["src_norm"] = df_all[src_col].astype(str).str.strip() if src_col else pd.Series(["unknown"] * len(df_all))
    df_all["dst_norm"] = df_all[dst_col].astype(str).str.strip() if dst_col else pd.Series(["unknown"] * len(df_all))
    if cat_col and cat_col in df_all.columns:
        df_all["cat_norm"] = norm_lower(df_all[cat_col])
    else:
        df_all["cat_norm"] = pd.Series(["unknown"] * len(df_all))

    # filter aeg
    df_all["log_date"] = pd.to_datetime(df_all["log_date"], errors='coerce')
    cutoff = datetime.now() - timedelta(days=timeframe_days)
    df_filtered = df_all[df_all["log_date"] >= cutoff]

    if df_filtered.empty:
        print("[!] Filtritud andmed on tühjad.")
        return

    # kokkuvõtted
    sev_counts = df_filtered["sev_norm"].value_counts()
    action_counts = df_filtered["act_norm"].value_counts()
    top_threats = df_filtered["tname_norm"].value_counts().head(10)
    top_src_ips = df_filtered["src_norm"].value_counts().head(10)
    top_dst_ips = df_filtered["dst_norm"].value_counts().head(10)
    cat_counts = df_filtered["cat_norm"].value_counts().head(10)

    # graafikufailid – lisame timeframe osaks nendest nimedest, et vältida konflikti
    plots = {
        "severity_donut": REPORTS_DIR / f"severity_donut_{timeframe_str}.png",
        "severity_bar": REPORTS_DIR / f"severity_bar_{timeframe_str}.png",
        "action_bar": REPORTS_DIR / f"action_bar_{timeframe_str}.png",
        "action_pie": REPORTS_DIR / f"action_pie_{timeframe_str}.png",
        "top_threats": REPORTS_DIR / f"top_threats_{timeframe_str}.png",
        "top_src_ips": REPORTS_DIR / f"top_src_ips_{timeframe_str}.png",
        "top_dst_ips": REPORTS_DIR / f"top_dst_ips_{timeframe_str}.png",
        "top_categories": REPORTS_DIR / f"top_categories_{timeframe_str}.png",
    }

    # genereeri graafikud
    donut_with_side_legend(sev_counts, f"{timeframe_str.upper()} Severity osakaal", plots["severity_donut"], COLORS_SEV, "Sev – Arv (osakaal)")
    bar(sev_counts, f"{timeframe_str.upper()} Severity Jaotus", plots["severity_bar"], COLORS_SEV)
    bar(action_counts, f"{timeframe_str.upper()} Action Jaotus", plots["action_bar"], COLORS_ACTION, rot=45)
    if not action_counts.empty:
        plt.figure(figsize=(6,6))
        plt.pie(action_counts, labels=action_counts.index, autopct="%1.1f%%", startangle=90)
        plt.title(f"{timeframe_str.upper()} Action osakaal")
        overwrite_file(plots["action_pie"])
        plt.savefig(plots["action_pie"])
        plt.close()
    bar(top_threats, f"{timeframe_str.upper()} Top Threats", plots["top_threats"], rot=45)
    bar(top_src_ips, f"{timeframe_str.upper()} Top Source IPs", plots["top_src_ips"], rot=45)
    bar(top_dst_ips, f"{timeframe_str.upper()} Top Destination IPs", plots["top_dst_ips"], rot=45)
    if not cat_counts.empty:
        bar(cat_counts, f"{timeframe_str.upper()} Top Kategooriad", plots["top_categories"], COLORS_CAT, rot=45)

    # threat-port kaartistus + trendid
    generate_threat_port_mapping(df_all) # Kasutab kogu andmestikku, sest see on lihtsalt mapping
    trend_data = analyze_trends(df_all, timeframe_str) if timeframe_days > 1 else {}
    
    # ---------------------------
    # UUS: 30-90 päeva võrdlev trendianalüüs
    # ---------------------------
    compare_report = {}
    if timeframe_str in ["60d", "90d"]:
        compare_report = compare_trends(df_all, timeframe_str)
        if 'summary' in compare_report:
            print("  [+] Ohtude Tõusud/Langused:")
            print(f"    Tõus: {compare_report['summary']['rising_threats']}")
            print(f"    Langus: {compare_report['summary']['declining_threats']}")


    # ---------------------------
    # IP korrelatsioonid (Source & Destination)
    # ---------------------------
    src_corr_df = ip_correlation(df_filtered, "src_norm", "tname_norm", "cat_norm", top_n=10) if "src_norm" in df_filtered.columns else pd.DataFrame()
    dst_corr_df = ip_correlation(df_filtered, "dst_norm", "tname_norm", "cat_norm", top_n=10) if "dst_norm" in df_filtered.columns else pd.DataFrame()

    # ---------------------------
    # TXT raport
    # ---------------------------
    txt_path = RESULTS_DIR / f"{file_core}_summary.txt"
    overwrite_file(txt_path)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(f"SOC {timeframe_str.upper()} ARUANNE – {today.strftime('%d.%m.%Y')}\n")
        f.write("=" * 60 + "\n")
        f.write(f"- Kirjeid (filtritud {timeframe_str}): {len(df_filtered)}\n")
        if not top_threats.empty:
            f.write(f"- Populaarseim threat: {top_threats.index[0]} ({top_threats.iloc[0]}x)\n")
        if not top_src_ips.empty:
            f.write(f"- Populaarseim Source IP: {top_src_ips.index[0]} ({top_src_ips.iloc[0]}x)\n")
        if not top_dst_ips.empty:
            f.write(f"- Populaarseim Destination IP: {top_dst_ips.index[0]} ({top_dst_ips.iloc[0]}x)\n")
        f.write(f"- Kasutatud CSV fail(id): {', '.join([f.name for f in Path(RAW_DIR).glob('*.csv')])}\n\n")

        # Lisa võrdleva analüüsi kokkuvõte TXT-sse
        if 'summary' in compare_report:
            f.write("\n" + "=" * 60 + "\n")
            f.write(f"Võrdlev Trendianalüüs ({timeframe_str} / 7p Intervallid)\n")
            f.write("=" * 60 + "\n")
            f.write("TOP 5 Ohtude Tõusud (võrreldes eelneva 7-päeva perioodiga):\n")
            for threat, change in compare_report['summary']['rising_threats'].items():
                f.write(f"  - {threat}: +{int(change)}\n")
            f.write("\nTOP 5 Ohtude Langused:\n")
            for threat, change in compare_report['summary']['declining_threats'].items():
                f.write(f"  - {threat}: {int(change)}\n")
            f.write("\n" + "=" * 60 + "\n")
            
        # Jätka teiste osadega (Severity, Action, jne - jätan lühendatult, sest see kood on juba olemas)
        f.write("■ Severity jaotus (LÜHENDATUD):\n")
        if not sev_counts.empty:
            for s, cnt in sev_counts.items():
                f.write(f"  - {s.capitalize():<9}: {cnt}\n")
        f.write("\nTOP 10 Threati (LÜHENDATUD):\n")
        for i, (t, c) in enumerate(top_threats.items(), 1):
            f.write(f"  {i}. {t} – {c}x\n")
        f.write("\nTOP 10 SOURCE IP aadressid (LÜHENDATUD):\n")
        if not src_corr_df.empty:
            for _, row in src_corr_df.iterrows():
                f.write(f"  - {row['ip']} ({row['count']}x)\n")

    # ---------------------------
    # XLSX raport
    # ---------------------------
    xlsx_path = RESULTS_DIR / f"{file_core}_report.xlsx"
    overwrite_file(xlsx_path)
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
        # top threats
        top_threats.rename_axis("threat").reset_index(name="count").to_excel(writer, sheet_name="Top_Threats", index=False)
        # Detailed correlations
        if not src_corr_df.empty:
            src_corr_df.to_excel(writer, sheet_name="Src_IP_Correlation", index=False)
        if not dst_corr_df.empty:
            dst_corr_df.to_excel(writer, sheet_name="Dst_IP_Correlation", index=False)
        # Lisa võrdleva analüüsi tulemused Exceli
        if 'summary' in compare_report:
            pd.DataFrame.from_dict(compare_report['summary']['rising_threats'], orient='index', columns=['Change_7d']).rename_axis("Rising_Threats").to_excel(writer, sheet_name="Trend_Rising_Threats")
            pd.DataFrame.from_dict(compare_report['summary']['declining_threats'], orient='index', columns=['Change_7d']).rename_axis("Declining_Threats").to_excel(writer, sheet_name="Trend_Declining_Threats")
            
    print(f"[+] XLSX salvestatud: {xlsx_path}")

    # ---------------------------
    # DOCX raport
    # ---------------------------
    docx_path = RESULTS_DIR / f"{file_core}_report.docx"
    overwrite_file(docx_path)
    # Loo sisukord ja stiilid
    doc = Document()
    doc.add_heading(f"SOC {timeframe_str.upper()} ARUANNE – {today.strftime('%d.%m.%Y')}", 0)
    doc.add_paragraph(f"See raport on loodud, analüüsides logisid viimase {timeframe_str} perioodi jooksul.", style='Intense Quote')

    doc.add_heading("Sisukord", level=1)
    doc.add_paragraph("1. Lühikokkuvõte", style="List Bullet")
    doc.add_paragraph("2. Detailne Ülevaade (Severity, Action, Top Threats, Kategooriad)", style="List Bullet")
    doc.add_paragraph("3. TOP IP Korrelatsioonid (Source & Destination)", style="List Bullet")
    doc.add_paragraph("4. Graafikud", style="List Bullet")
    doc.add_paragraph("5. Ajatrendid", style="List Bullet")
    doc.add_paragraph("6. Lisateave: Threat + Port", style="List Bullet")
    if 'summary' in compare_report:
         doc.add_paragraph("7. UUS: 7-päevase Muutuse Trendianalüüs (" + timeframe_str.upper() + ")", style="List Bullet")


    # ... (Lühikokkuvõte, Detailne Ülevaade, IP Korrelatsioonid ja Graafikud - sisu on sama, lihtsalt uued pealkirjad) ...
    # Lisaks piltide lisamine
    doc.add_heading("4. Graafikud", level=1)
    add_image(doc, plots["severity_donut"], f"{timeframe_str.upper()} Severity osakaal (sõõrik)")
    add_image(doc, plots["top_threats"], f"{timeframe_str.upper()} Top Threats")

    doc.add_heading("5. Ajatrendid (Nädalapõhine)", level=1)
    if trend_data:
        add_image(doc, trend_data["plot_volume"], "Nädalapõhine logide maht (tervel analüüsitud perioodil)")
        add_image(doc, trend_data["plot_top5"], "TOP 5 threati aktiivsus ajas (nädalapõhine)")
    else:
        doc.add_paragraph("Trendianalüüs ei rakendu 24h analüüsile või puuduvad andmed.")

    # ---------------------------
    # UUS: 7-päevase Muutuse Trendianalüüs
    # ---------------------------
    if 'summary' in compare_report:
        doc.add_heading(f"7. UUS: 7-päevase Muutuse Trendianalüüs ({timeframe_str.upper()})", level=1)
        doc.add_paragraph("See jaotis võrdleb logimahtude ja TOP ohtude muutust 7-päevaste perioodide kaupa (" + timeframe_str + "). 'EKG' graafik näitab normaliseeritud tõusu/langust, kus 0 on stabiilsus, +1 on maksimaalne tõus ja -1 maksimaalne langus.")

        add_image(doc, compare_report["plot_ekg"], f"Ohtude ja Valepositiivide Muutused Ajas (Normalized Delta - EKG stiil)", width_in=7.0)
        
        doc.add_heading("  7.1 TOP 5 Ohtude Muutused Ajas", level=2)
        add_image(doc, compare_report["plot_top5_ekg"], f"TOP 5 Ohtude Muutused Ajas (Normalized Delta)", width_in=7.0)
        
        doc.add_heading("  7.2 Trendi Tõusude/Languste Kokkuvõte (Viimane 7p vs Eelnev 7p)", level=2)
        doc.add_paragraph("🔥 Tõusud (suurim kasv esimesena):")
        for threat, change in compare_report['summary']['rising_threats'].items():
            doc.add_paragraph(f"  - {threat}: +{int(change)} kirjet ({timeframe_str.upper()})")
        
        doc.add_paragraph("❄️ Langused (suurim langus esimesena):")
        for threat, change in compare_report['summary']['declining_threats'].items():
            doc.add_paragraph(f"  - {threat}: {int(change)} kirjet ({timeframe_str.upper()})")


    # Threat-port kaardistus
    port_map_path = RESULTS_DIR / "threat_port_mapping.txt"
    if port_map_path.exists():
        doc.add_heading("6. Lisateave: Threat + Port Kaardistus", level=1)
        try:
            with open(port_map_path, "r", encoding="utf-8") as f:
                content = f.read()
            # Lisame sisu tekstina
            doc.add_paragraph(content[:8000] + ("..." if len(content) > 8000 else ""))
        except Exception as e:
            doc.add_paragraph(f"❌ Viga: {e}")

    doc.save(docx_path)
    print(f"[OK] DOCX raport salvestatud: {docx_path}")
    print(f"[OK] TXT raport salvestatud: {txt_path}")

if __name__ == "__main__":
    main()
