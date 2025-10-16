# ==============================================================
#  SOC 24h logianalüüs — v3.0
#  Autor: ocrHeiki (GitHub: https://github.com/ocrHeiki)
#  Kirjeldus:
#   - Analüüsib uusimat CSV logifaili kaustas raw/
#   - Loob TXT, DOCX, CSV ja XLSX aruanded
#   - Taastatud TOP 10 Threat / Content Name graafik ja tekstiosa
#   - Lisatud 24h_ip.txt – lähte-IP nimekiri edasisteks analüüsideks
# ==============================================================

from pathlib import Path
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Kaustad ---
BASE = Path.home() / "Documents" / "SOC"
RAW = BASE / "raw"
OUT = BASE / "tulemused"
REP = BASE / "reports"
for d in (RAW, OUT, REP):
    d.mkdir(parents=True, exist_ok=True)

# --- Värvikaardid ---
COLORS_SEV = {"low": "#0000FF", "medium": "#FFFF00", "high": "#FFA500", "critical": "#FF0000"}
COLORS_ACTION = {"allow": "#33CC33", "deny": "#CC3333", "drop": "#3366CC", "alert": "#FFCC00", "reset-both": "#9933CC", "reset-server": "#800080"}
COLORS_TYPE = {"malware": "#CC0033", "vulnerability": "#FF3333", "spyware": "#3399FF", "suspicious": "#FFCC66", "benign": "#66CC66"}
COLORS_CAT = {
    "command-and-control": "#CC0000", "code-execution": "#FF6600", "sql-injection": "#FF9933",
    "brute-force": "#FFCC00", "dos": "#FFFF66", "hacktool": "#9933CC", "info-leak": "#66CCFF",
    "spyware": "#3399FF", "code-obfuscation": "#996633"
}

# --- Abifunktsioonid ---
def iso_from_filename(name: str):
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", name)
    if m:
        return f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
    m2 = re.search(r"(\d{4})-(\d{2})-(\d{2})", name)
    return m2.group(0) if m2 else datetime.now().date().isoformat()

def first_existing(df, names):
    for n in names:
        if n in df.columns:
            return n
    return None

def norm_lower(s: pd.Series) -> pd.Series:
    return s.astype(str).str.strip().str.lower()

def bar(series, title, outpath, colors=None, rot=0):
    if series is None or series.empty:
        return
    plt.figure(figsize=(10, 5))
    c = [colors.get(str(i).lower(), "#888888") for i in series.index] if colors else "#888888"
    series.plot(kind="bar", color=c)
    plt.title(title)
    plt.xticks(rotation=rot, ha="right" if rot else "center")
    plt.tight_layout()
    plt.savefig(outpath)
    plt.close()

def pie(series, title, outpath):
    if series is None or series.empty:
        return
    plt.figure(figsize=(6, 6))
    plt.pie(series, labels=series.index, autopct="%1.1f%%", startangle=90)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(outpath)
    plt.close()

def add_image(doc: Document, img_path: Path, caption: str, width_in=6.0):
    if img_path.exists():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_in))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- Peamine funktsioon ---
def main():
    csv_files = sorted(RAW.glob("*.csv"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not csv_files:
        print("[!] Ei leitud CSV-faile kaustas 'raw/'.")
        return

    path = csv_files[0]
    date_iso = iso_from_filename(path.name)
    print(f"[i] Analüüsitakse (uusim): {path.name} ({date_iso})")

    df = pd.read_csv(path, encoding="utf-8", low_memory=False)

    # Veerud
    sev_col = first_existing(df, ["Severity", "severity"])
    act_col = first_existing(df, ["Action", "action"])
    cat_col = first_existing(df, ["thr_category", "category", "Threat Category"])
    type_col = first_existing(df, ["Threat/Content Type", "Threat Type", "threat_type"])
    name_col = first_existing(df, ["Threat/Content Name", "Threat Name", "content_name", "threat_name"])
    src_col = first_existing(df, ["Source address", "Source", "src", "src_ip"])
    dst_col = first_existing(df, ["Destination address", "Destination", "dst", "dst_ip"])

    if sev_col: df["sev_norm"] = norm_lower(df[sev_col])
    if act_col: df["act_norm"] = norm_lower(df[act_col])
    if cat_col: df["cat_norm"] = norm_lower(df[cat_col])
    if type_col: df["type_norm"] = norm_lower(df[type_col])
    if name_col: df["tname_norm"] = df[name_col].astype(str).str.strip()
    if src_col: df["src_norm"] = df[src_col].astype(str).str.strip()
    if dst_col: df["dst_norm"] = df[dst_col].astype(str).str.strip()

    # Statistika
    total = len(df)
    sev_counts = df["sev_norm"].value_counts() if "sev_norm" in df.columns else pd.Series(dtype=int)
    act_counts = df["act_norm"].value_counts() if "act_norm" in df.columns else pd.Series(dtype=int)
    top_cat = df["cat_norm"].value_counts().head(10) if "cat_norm" in df.columns else pd.Series(dtype=int)
    top_src = df["src_norm"].value_counts().head(10) if "src_norm" in df.columns else pd.Series(dtype=int)
    top_dst = df["dst_norm"].value_counts().head(10) if "dst_norm" in df.columns else pd.Series(dtype=int)
    top_threat = df["tname_norm"].value_counts().head(10) if "tname_norm" in df.columns else pd.Series(dtype=int)

    # --- TXT raport ---
    out_txt = OUT / f"24h_summary_{date_iso}.txt"
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write(f"SOC 24h KOONDARUANNE – {date_iso}\n")
        f.write("=" * 50 + "\n")
        f.write(f"Logifail: {path.name}\nKokku ridu: {total}\n\n")
        if not sev_counts.empty:
            f.write("■ Severity jaotus:\n")
            for s, c in sev_counts.items():
                f.write(f"  - {s.title():<10}: {c}\n")
        if not act_counts.empty:
            f.write("\n■ Action jaotus:\n")
            for a, c in act_counts.items():
                f.write(f"  - {a:<10}: {c}\n")
        if not top_cat.empty:
            f.write("\n■ TOP kategooriad:\n")
            for i, (k, v) in enumerate(top_cat.items(), 1):
                f.write(f"  {i}. {k} – {v}\n")
        if not top_threat.empty:
            f.write("\n■ TOP 10 Threat / Content Name:\n")
            for i, (k, v) in enumerate(top_threat.items(), 1):
                f.write(f"  {i}. {k} – {v}\n")
        if not top_src.empty:
            f.write("\n■ TOP 10 allika IP:\n")
            for i, (k, v) in enumerate(top_src.items(), 1):
                f.write(f"  {i}. {k} – {v}\n")
        if not top_dst.empty:
            f.write("\n■ TOP 10 sihtmärgi IP:\n")
            for i, (k, v) in enumerate(top_dst.items(), 1):
                f.write(f"  {i}. {k} – {v}\n")

    # --- IP nimekiri edasisteks analüüsideks ---
    ip_path = OUT / f"24h_ip_{date_iso}.txt"
    with open(ip_path, "w", encoding="utf-8") as f:
        f.write("TOP 10 Allika IP (24h)\n")
        f.write("=" * 30 + "\n\n")
        for ip, cnt in top_src.items():
            f.write(f"{ip} ({cnt} korda)\n")

    # --- XLSX ---
    out_xlsx = OUT / f"24h_summary_{date_iso}.xlsx"
    with pd.ExcelWriter(out_xlsx) as xw:
        info_df = pd.DataFrame({
            "Key": ["Logifail", "Kuupäev", "Kirjeid kokku", "Analüüsi aeg"],
            "Value": [path.name, date_iso, total, datetime.now().strftime("%Y-%m-%d %H:%M:%S")]
        })
        info_df.to_excel(xw, sheet_name="Info", index=False)
        if not sev_counts.empty: sev_counts.rename_axis("Severity").reset_index(name="Count").to_excel(xw, sheet_name="Severity", index=False)
        if not act_counts.empty: act_counts.rename_axis("Action").reset_index(name="Count").to_excel(xw, sheet_name="Action", index=False)
        if not top_cat.empty: top_cat.rename_axis("Category").reset_index(name="Count").to_excel(xw, sheet_name="TopCategories", index=False)
        if not top_threat.empty: top_threat.rename_axis("ThreatName").reset_index(name="Count").to_excel(xw, sheet_name="TopThreats", index=False)
        if not top_src.empty: top_src.rename_axis("SrcIP").reset_index(name="Count").to_excel(xw, sheet_name="TopSrc", index=False)
        if not top_dst.empty: top_dst.rename_axis("DstIP").reset_index(name="Count").to_excel(xw, sheet_name="TopDst", index=False)

    # --- Graafikud ---
    bar(sev_counts, f"■ Severity jaotus (24h) – {date_iso}", REP / f"paev_severity_bar_{date_iso}.png", colors=COLORS_SEV)
    pie(sev_counts, f"▲ Severity osakaal (24h) – {date_iso}", REP / f"paev_severity_pie_{date_iso}.png")
    bar(act_counts, f"■ Action jaotus (24h) – {date_iso}", REP / f"paev_action_bar_{date_iso}.png", colors=COLORS_ACTION, rot=45)
    pie(act_counts, f"▲ Action osakaal (24h) – {date_iso}", REP / f"paev_action_pie_{date_iso}.png")
    bar(top_cat, f"■ TOP kategooriad (24h) – {date_iso}", REP / f"paev_top_cat_{date_iso}.png", colors=COLORS_CAT, rot=45)
    bar(top_threat, f"■ TOP 10 Threat / Content Name (24h) – {date_iso}", REP / f"paev_top_threats_{date_iso}.png", rot=45)
    bar(top_src, f"■ TOP 10 allika IP (24h) – {date_iso}", REP / f"paev_top_src_{date_iso}.png", rot=45)
    bar(top_dst, f"■ TOP 10 sihtmärgi IP (24h) – {date_iso}", REP / f"paev_top_dst_{date_iso}.png", rot=45)

    # --- DOCX ---
    doc = Document()
    doc.add_heading(f"SOC 24h aruanne — {date_iso}", level=1)
    doc.add_heading("Tekstiline kokkuvõte", level=2)
    with open(out_txt, "r", encoding="utf-8") as fh:
        for line in fh:
            doc.add_paragraph(line.rstrip("\n"))

    doc.add_heading("Graafikud ja visuaalid", level=2)
    for img, cap in [
        (REP / f"paev_severity_bar_{date_iso}.png", "Severity jaotus (tulpdiagramm)"),
        (REP / f"paev_severity_pie_{date_iso}.png", "Severity osakaal (pirukas)"),
        (REP / f"paev_action_bar_{date_iso}.png", "Action jaotus (tulpdiagramm)"),
        (REP / f"paev_action_pie_{date_iso}.png", "Action osakaal (pirukas)"),
        (REP / f"paev_top_cat_{date_iso}.png", "TOP kategooriad"),
        (REP / f"paev_top_threats_{date_iso}.png", "TOP Threat / Content Name"),
        (REP / f"paev_top_src_{date_iso}.png", "TOP allika IP"),
        (REP / f"paev_top_dst_{date_iso}.png", "TOP sihtmärgi IP")
    ]:
        add_image(doc, img, cap)

    docx_path = OUT / f"24h_summary_{date_iso}.docx"
    doc.save(str(docx_path))

    # --- Teated ---
    print("\n[OK] 24h analüüs valmis.")
    print(f" - TXT : {out_txt}")
    print(f" - XLSX: {out_xlsx}")
    print(f" - DOCX: {docx_path}")
    print(f" - IP nimekiri: {ip_path}")
    print(f" - Graafikud: {REP}")

if __name__ == "__main__":
    main()
