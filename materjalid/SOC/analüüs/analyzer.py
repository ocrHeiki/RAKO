# ==============================================================
#  SOC Threat Analyser v5.1
#  Autor: Heiki Rebane (õpiprojekt)
#  Kuupäev: 17.11.2025
#  Kirjeldus:
#   - Loeb CSV-failid kaustast raw/ ja filtreerib ajavahemiku järgi.
#   - Analüüsib threat-nimesid, IP aktiivsust ja sündmuste jagunemist.
#   - Lisab WEEKLY TREND ANALYSIS mooduli.
#   - Tuvastab Volume-Based ja CVE Diversity IP-d.
#   - Lisab MITRE ATT&CK kaardistuse ja Threat Vault info (koos cache'iga).
#   - Koostab DOCX, TXT ja XLSX raportid graafikutega.
#   - Kasutab MOCK GeoIP lahendust (pole vaja lisafaile alla laadida).
# ==============================================================

import subprocess
import sys
import os
import importlib
from pathlib import Path
from datetime import datetime, timedelta
import pandas as pd
import matplotlib.pyplot as plt
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import requests
import json
import matplotlib.dates as mdates # Kasutatakse trend graafikute kuupäevade kuvamiseks
import io # Kasutatakse XLSX-i kirjutamiseks

# ————————————————————————————————
# CLI Argumentide seadistus.
# Võimaldab määrata analüüsi ajavahemiku otse käsurealt.
# ————————————————————————————————

def parse_args():
    parser = argparse.ArgumentParser(description="SOC Threat Analyser - Palo Alto logianalüüs")
    parser.add_argument("--timeframe", choices=["24h", "7d", "30d"], default="7d",
                        help="Analüüsitav ajavahemik (vaikimisi: '7d')")
    parser.add_argument("--output", choices=["brief", "detailed"], default="detailed",
                        help="Aruande detailne tase (vaikimisi: 'detailed')")
    parser.add_argument("--strict-local", action="store_true",
                        help="Kohalike (France/Réunion) IP-de kohta rakendatakse karmimaid reegleid")
    return parser.parse_args()


# ————————————————————————————————
# Algne projekti struktuur ja kaustade seaded
# Määratletakse kõik sisend- ja väljundkaustad.
# ————————————————————————————————

BASE_DIR = Path(__file__).resolve().parent.parent
RAW_DIR = BASE_DIR / "raw" # Logifailid
REPORTS_DIR = BASE_DIR / "reports" # Graafikud
RESULTS_DIR = BASE_DIR / "tulemused" # Lõppraportid
THREAT_VAULT_CACHE = BASE_DIR / "threat_vault_cache" # Threat Vault andmete vahemälu
TRENDS_DIR = BASE_DIR / "trendid" # Kaust trendiraportite graafikute jaoks

# Loob kõik vajalikud kaustad, kui need puuduvad
for d in [RAW_DIR, REPORTS_DIR, RESULTS_DIR, THREAT_VAULT_CACHE, TRENDS_DIR]:
    d.mkdir(parents=True, exist_ok=True) 

# Seadistused ja püsivad andmed
COLORS_SEV = {"low": "#0000FF", "medium": "#FFFF00", "high": "#FFA500", "critical": "#FF0000"}

# Riskianalüüs valesid nähtusi
fp_guidance = {
    "Nmap Aggressive Option Print Detection": {
        "risk": "KÕRGE",
        "reason": "Sageli kasutavad seda legitiimsed süsteemihaldurid.",
        "tip": "Kontrolli IP konteksti. Võib viidata pentestile või turvakontrollile."
    }
}

# MITRE ATT&CK kaardistus threatimede järgi
attck_mapping = {
    "Nmap Aggressive Option Print Detection": {"tactic": "Discovery", "technique": "T1046"},
    "HTTP2 Protocol Suspicious RST STREAM Frame detection": {"tactic": "Defense Evasion", "technique": "T1071"},
    "Windows Local Security Authority lsardelete access": {"tactic": "Credential Access", "technique": "T1003"}
}

# MOCK GeoIP – Kasutatakse ainult juhul, kui päris GeoIP2 faile pole lubatud alla laadida.
predefined_geo = {
    "192.168.1.20": {"country": "Prantsusmaa", "city": "Pariis"},
    "192.168.2.100": {"country": "Réunion", "city": "Saint-Denis"},
    "8.8.8.8": {"country": "Ameerika Ühendriigid", "city": "Kalifornia"},
    "203.12.160.45": {"country": "Hiina", "city": "Shenzhen"},
    "176.10.10.1": {"country": "Venemaa", "city": "Moskva"}
}

def get_country(ip):
    # Tagastab IP-l vastava riigi nime või "Teadmata"
    return predefined_geo.get(ip, {}).get("country", "Teadmata")

# ————————————————————————————————
# Threat Vault integratsioon (sisseehitatud API päring ja cache)
# ————————————————————————————————

def get_threat_details(threat_name):
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', threat_name)
    cache_file = THREAT_VAULT_CACHE / f"{safe_name}.json"
    
    # 1. Kontrolli cache'i olemasolu: väldib korduvaid võrgupäringuid
    if cache_file.exists():
        try:
            with open(cache_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except json.JSONDecodeError:
             print(f"❗Viga cache'i lugemisel failist: {cache_file}")

    # 2. Tee reaalajas päring Palo Alto Threat Vault API-le
    try:
        url = "https://threatvault.paloaltonetworks.com/restapi/threats"
        headers = {"User-Agent": "SOC-Threat-Analyser"}
        params = {"search": threat_name}
        resp = requests.get(url, headers=headers, params=params, timeout=10) # 10 sekundi timeout
        
        if resp.status_code == 200:
            data = resp.json()
            threat = data.get("threats", [])[0] if data.get("threats") else {}
            
            # 3. Salvesta vastus cache'i faili
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(threat, f, indent=2)
            return threat
    
    except requests.exceptions.RequestException as e:
        print(f"❗Võrgu/API viga Threat Vault otsingus '{threat_name}': {e}")
    except Exception as e:
        print(f"❗Muu viga Threat Vault otsingus '{threat_name}': {e}")
    
    return {} # Tagastab tühja, kui päring ebaõnnestus

# ————————————————————————————————
# Abifunktsioonid
# ————————————————————————————————

def iso_from_filename(name: str):
    # Ekstraktib kuupäeva failinimest (nt logi_17.11.2025.csv)
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", name)
    if m:
        return datetime.strptime(f"{m.group(3)}-{m.group(2)}-{m.group(1)}", "%Y-%m-%d").date()
    m = re.search(r"(\d{4})-(\d{2})-(\d{2})", name)
    return m.group(0) if m else datetime.now().date()

def first_existing(df, names):
    # Leiab logifailist esimese sobiva veeru nime (nt 'Severity' või 'Risk')
    for n in names:
        if n in df.columns:
            return n
    return None

def norm_lower(col):
    # Normaliseerib andmed: kõik väiketähtedeks, tühikud ja reavahetused eemaldatud
    return col.astype(str).str.lower().str.strip()

def bar(series, title, outpath, colors=None, rot=0, is_date_trend=False):
    # Loob ja salvestab tulpdiagrammi
    if series.empty: return
    
    plt.figure(figsize=(10, 5))
    ax = series.plot(kind="bar", color="#3399FF")
    plt.title(title)
    
    if is_date_trend:
        # Eraldi käsitlus trendi graafikute jaoks (kuupäevadega)
        ax.set_xticklabels(series.index.strftime('%Y-%m-%d'))
        plt.xticks(rotation=45, ha='right')
    else:
        plt.xticks(rotation=rot)
        
    plt.tight_layout()
    plt.savefig(outpath)
    plt.close()

def add_image(doc, img_path, caption, width_in=6.0):
    # Lisab graafiku Microsoft Word (DOCX) raportisse
    if not img_path.exists(): return
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ————————————————————————————————
# TRENDIANALÜÜSI MOODUL
# Analüüsib tegevuse mahtu ja top threatide arengut aja jooksul.
# ————————————————————————————————

def analyze_trends(df_all):
    print("📈 Käivitan nädalapõhise trendianalüüsi...")

    # Teeb andmetest koopia, et mitte muuta originaal DataFrame'i
    df_trends = df_all.copy()
    
    # Kindlustab, et log_date on kuupäeva formaadis ja ilma NaN/NaT väärtusteta
    df_trends["log_date"] = pd.to_datetime(df_trends["log_date"], errors='coerce') 
    df_trends.dropna(subset=["log_date"], inplace=True)

    # Arvutab iga logi kirje jaoks vastava nädala alguse kuupäeva (grupeerimiseks)
    df_trends['week_start'] = df_trends['log_date'].apply(
        lambda x: x - timedelta(days=x.weekday()) # Leiab esmaspäeva
    ).dt.normalize()

    # 1. Kokkuvõtlik rünnete maht nädalate kaupa (Volume)
    # Grupeerib andmed nädala alguse järgi ja loendab kirjed
    weekly_volume = df_trends.groupby('week_start').size() 
    
    # Salvestab mahu graafiku
    out_trend_vol = TRENDS_DIR / "trend_weekly_volume.png"
    bar(weekly_volume, "Nädalapõhine Logide Maht (Trend)", out_trend_vol, is_date_trend=True)

    # 2. TOP 5 Threatide trendid ajas (Aktiivsus)
    # Tuvastab 5 kõige levinumat ohunime kogu andmekogumis
    top5_threat_names = df_trends['tname_norm'].value_counts().head(5).index.tolist()
    
    # Grupeeri nädala ja threati nime järgi, et näha aktiivsust top5 threati jaoks igas nädalas
    weekly_threats = df_trends.groupby(['week_start', 'tname_norm']).size().unstack(fill_value=0)
    
    # Valib ainult Top 5 nimega seotud veerud
    top5_trend = weekly_threats[weekly_threats.columns.intersection(top5_threat_names)]

    # Salvestab Top 5 threatide graafiku
    out_trend_top5 = TRENDS_DIR / "trend_top5_threats.png"
    plt.figure(figsize=(12, 6))
    top5_trend.plot(kind='line', marker='o', ax=plt.gca()) 
    plt.title("TOP 5 Threat'i Aktiivsus Aja Jooksul")
    plt.xlabel("Nädala Algus")
    plt.ylabel("Kirjete Arv")
    plt.legend(title='Threat', bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(out_trend_top5)
    plt.close()
    
    # Koostab tulemuste dictionary
    trend_result = {
        "weekly_volume": weekly_volume.to_string(), 
        "top5_trend": top5_trend.to_string(),
        "plot_volume": out_trend_vol,
        "plot_top5": out_trend_top5
    }
    
    print("✔️ Trendianalüüs lõpetatud. Tulemused salvestatud /trendid kausta.")
    return trend_result


# ————————————————————————————————
# PROJEKTI PÕHIPLOKK – ANALÜÜSI JA RAPORTITE LOOMINE
# Siin toimub andmete pärimine, töötlemine ja aruandlus.
# ————————————————————————————————

def main():
    args = parse_args()
    # Teisendab ajavahemiku päevade arvuks
    timeframe_days = {"24h": 1, "7d": 7, "30d": 30}[args.timeframe]

    print(f"🔍 Käivitan {args.timeframe} analüüsi ({timeframe_days} päeva) | Strict Local: {args.strict_local}")

    csv_files = list(RAW_DIR.glob("*.csv"))
    if not csv_files:
        print("⚠️ Ühtegi CSV-faili ei leitud kaustas: raw/")
        return

    dfs = []
    today = datetime.now().date()

    # Loeb kõik CSV failid ühte nimekirja
    for f in csv_files:
        df = pd.read_csv(f, low_memory=False)
        date = iso_from_filename(f.name)
        df["log_date"] = date
        dfs.append(df)

    # Kombineerib kõik nimekirjas olevad DataFrame'id üheks suureks
    df_all = pd.concat(dfs, ignore_index=True)
    print(f"✔️ Loetud {len(df_all)} kirjet {len(csv_files)} logifailist")

    # Ajavahemiku filtreerimine
    df_all["log_date"] = pd.to_datetime(df_all["log_date"])
    cutoff_date = today - timedelta(days=timeframe_days)
    df_filtered = df_all[df_all["log_date"].dt.date >= cutoff_date]
    
    # ----------------------------------------------------
    # Kutsutakse trendi analüüsi funktsioon väärtustega, mis lähevad aruandesse
    trend_data = analyze_trends(df_all)
    # ----------------------------------------------------


    # ————————————————————————————————
    # VEERGUDE TUVASTAMINE JA NORMALISEERIMINE
    # Tagab, et saame logisid töödelda sõltumata veergude nime erinevustest.
    # ————————————————————————————————
    sev_col = first_existing(df_filtered, ["Severity", "severity", "risk"])
    act_col = first_existing(df_filtered, ["Action", "action"])
    threat_col = first_existing(df_filtered, ["Threat/Content Name", "threat_name"])
    src_col = first_existing(df_filtered, ["Source address", "src", "source"])
    dst_col = first_existing(df_filtered, ["Destination address", "dst"])
    rule_col = first_existing(df_filtered, ["Rule", "rule", "policy"])
    sport_col = first_existing(df_filtered, ["Source Port", "sport"])
    dport_col = first_existing(df_filtered, ["Destination Port", "dport"])
    time_col = first_existing(df_filtered, ["Receive Time", "time", "receive_time"])

    # Normaliseerib võtmeks olevad veerud (väiketähtedeks, tühikud eemaldatud)
    df_filtered["sev_norm"] = norm_lower(df_filtered[sev_col])
    df_filtered["act_norm"] = norm_lower(df_filtered[act_col])
    df_filtered["tname_norm"] = norm_lower(df_filtered[threat_col])
    df_filtered["src_norm"] = norm_lower(df_filtered[src_col])
    df_filtered["dst_norm"] = norm_lower(df_filtered[dst_col])

    # Lisab MITRE ja GeoIP infor
    df_filtered["mitre_tactic"] = df_filtered["tname_norm"].map(lambda x: attck_mapping.get(str(x), {}).get("tactic", "–"))
    df_filtered["mitre_technique"] = df_filtered["tname_norm"].map(lambda x: attck_mapping.get(str(x), {}).get("technique", "–"))
    df_filtered["src_country"] = df_filtered["src_norm"].apply(get_country) # Kasutab MOCK GeoIP

    # ————————————————————————————————
    # STATISTILISTE ANDMETE ARVUTAMINE
    # ————————————————————————————————
    sev_counts = df_filtered["sev_norm"].value_counts()
    action_counts = df_filtered["act_norm"].value_counts()
    top_threat = df_filtered["tname_norm"].value_counts().head(10)
    top_countries = df_filtered["src_country"].value_counts().head(10)


    # LOW Severity detailanalüüs (Vajalik detailsema reegliinfo tuvastamiseks)
    df_low = df_filtered[df_filtered["sev_norm"] == "low"]
    df_low_summary = pd.DataFrame()

    if not df_low.empty:
      low_threat_counts = df_low["tname_norm"].value_counts().head(10)
      low_threat_summary = []
      
      for threat in low_threat_counts.index:
            subset = df_low[df_low["tname_norm"] == threat]
            # Leiab kõige sagedamini kasutatud reegli, pordid jne.
            rule = subset[rule_col].mode().iloc[0] if rule_col and not subset[rule_col].mode().empty else "N/A"
            port_src = subset[sport_col].mode().iloc[0] if sport_col and not subset[sport_col].mode().empty else "N/A"
            port_dst = subset[dport_col].mode().iloc[0] if dport_col and not subset[dport_col].mode().empty else "N/A"
            unique_ips = subset["src_norm"].nunique()
            total = len(subset)

            low_threat_summary.append({
                "Threat Name": threat,
                "Kokku": total,
                "Unikaalsed IP-d": unique_ips,
                "Rule": rule,
                "Source Port": port_src,
                "Destination Port": port_dst
            })
            
      df_low_summary = pd.DataFrame(low_threat_summary)

      # Mõõdab kellaajalist aktiivsust (LOW threatide puhul)
      if time_col in df_low.columns:
          df_low[time_col] = pd.to_datetime(df_low[time_col], errors='coerce')
          df_low = df_low.dropna(subset=[time_col])
          df_low["hour"] = df_low[time_col].dt.hour
          hourly_activity = df_low["hour"].value_counts().sort_index()
      else:
          hourly_activity = pd.Series(dtype=int)
    else:
        hourly_activity = pd.Series(dtype=int)


    # Volume-Based Detection
    volume_threshold = 50
    volume_ips = df_filtered["src_norm"].value_counts()
    volume_ips = volume_ips[volume_ips > volume_threshold]
    # ... (df_volume loomine)
    df_volume = pd.DataFrame() # Asenda päris koodiga

    # CVE Diversity Detection
    cve_diversity_threshold = 10
    df_threat_per_ip = df_filtered.groupby("src_norm")["tname_norm"].nunique().sort_values(ascending=False)
    diverse_ips = df_threat_per_ip[df_threat_per_ip >= cve_diversity_threshold]
    # ... (df_cve loomine)
    df_cve = pd.DataFrame() # Asenda päris koodiga

    # Väljundfailid ja graafikute loomine
    time_range = args.timeframe
    today_str = today.strftime("%Y-%m-%d")
    out_txt = RESULTS_DIR / f"SOC_{time_range}_aruande_{today_str}.txt"
    out_docx = RESULTS_DIR / f"SOC_{time_range}_aruande_{today_str}.docx"
    
    # Graafikute salvestamine
    plot_hourly = REPORTS_DIR / f"hourly_activity_{time_range}_{today_str}.png"
    plot_top_threat = REPORTS_DIR / f"top_threats_{time_range}_{today_str}.png"

    if not hourly_activity.empty:
        bar(hourly_activity, "Kellaajaline aktiivsus (LOW Threatid)", plot_hourly, rot=0)
    if not top_threat.empty:
        bar(top_threat, "TOP 10 Threat Name", plot_top_threat, rot=45)


    # ————————————————————————————————
    # TXT RAPORTI GENERATSIOON
    # Masinloetav tekstiline väljund andmetest
    # ————————————————————————————————
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write(f"=== SOC {time_range.upper()} ARUANNE – {today_str} ===\n")
        f.write("Loetud logikirjete arv: {}\n".format(len(df_filtered)))
        f.write("\n.SEVERITY JAGUNEMINE:\n")
        f.write(sev_counts.to_string())
        
        # TREND TXT lisamine
        if trend_data:
            f.write("\n\n📊 NÄDALAPÕHINE TRENDIANALÜÜS:\n")
            f.write("\nVOLUME: Logimaht nädalate kaupa:\n")
            f.write(trend_data['weekly_volume'])
            f.write("\n\nTOP 5 THREAT: Aktiivsus nädalate kaupa:\n")
            f.write(trend_data['top5_trend'])
        
        # ... (Jätk TXT sisule)

    print(f"📝 TXT aruanne salvestatud: {out_txt}")


    # ————————————————————————————————
    # DOCX RAPORTI GENERATSIOON
    # Vormindatud väljund koos graafikute ja tabelitega
    # ————————————————————————————————
    doc = Document()
    doc.add_heading(f"SOC {time_range.upper()} ARUANNE – {today_str}", 0)
    
    # ... (Severity ja Riikide jaotus)
    
    # Lisame üldised graafikud
    if plot_top_threat.exists():
        add_image(doc, plot_top_threat, "TOP 10 Threat Name graafik", width_in=6.5)
    if plot_hourly.exists():
        add_image(doc, plot_hourly, "Kellaajaline aktiivsus (LOW Threatid)", width_in=6.5)

    # DOCX Trendide lisamine
    doc.add_page_break()
    doc.add_heading("Trendianalüüs", level=1)
    
    if trend_data:
        doc.add_paragraph("Volume: Logimaht nädalate kaupa:")
        doc.add_paragraph(trend_data['weekly_volume'])
        
        add_image(doc, trend_data['plot_volume'], "Nädala Logide Maht (Trend)", width_in=6.5)
        add_image(doc, trend_data['plot_top5'], "TOP 5 Threat'i Aktiivsus Aja Jooksul", width_in=6.5)
        
        doc.add_paragraph("TOP 5 Threat: Aktiivsus nädalate kaupa:")
        doc.add_paragraph(trend_data['top5_trend'])
        
    doc.add_page_break()
    doc.add_heading("TOP 10 Threat Name + Riskianalüüs", level=1)
    for i, (threat, count) in enumerate(top_threat.items()):
        if i >= 10: break
        doc.add_paragraph(f"{i+1}. {threat} – {count} korda")

        # MITRE Info
        tactic = attck_mapping.get(threat, {}).get("tactic", "–")
        technique = attck_mapping.get(threat, {}).get("technique", "–")
        doc.add_paragraph(f"🔗 MITRE: {tactic} | {technique}")

        # Threat Vault Info
        vault_info = get_threat_details(threat)
        if vault_info:
            doc.add_paragraph(f"🔍 Threat Vault ID: {vault_info.get('id', 'N/A')}")
            # ... (muu Threat Vault info)
        else:
            doc.add_paragraph("🔸 Threat Vault info puudub")
            
        # ... (Riskianalüüs)

    # ... (LOW, Volume ja CVE tabelid)
    doc.save(out_docx)
    print(f"📄 DOCX aruanne salvestatud: {out_docx}")

    # ————————————————————————————————
    # XLSX RAPORTI GENERATSIOON
    # Andmete eksport Excelisse detailsemaks töötluseks
    # ————————————————————————————————
    out_xlsx = RESULTS_DIR / f"SOC_{time_range}_aruande_{today_str}.xlsx"
    with pd.ExcelWriter(out_xlsx) as writer:
        sev_counts.to_frame("Arv").to_excel(writer, sheet_name="Severity")
        # ... (Muud üldised sheet'id)
        
        # XLSX Trendide lisamine
        if trend_data:
            try:
                # Konverteerib trendi stringid tagasi DataFrame'iks, et need Excelisse korrektselt kirjutada
                pd.read_csv(io.StringIO(trend_data['weekly_volume']), delim_whitespace=True).to_excel(writer, sheet_name="Trend-Volume", index=False)
                pd.read_csv(io.StringIO(trend_data['top5_trend']), delim_whitespace=True).to_excel(writer, sheet_name="Trend-Top5", index=False)
            except Exception as e:
                print(f"❗Viga trendiandmete XLSX-i kirjutamisel: {e}")

    print(f"📊 XLSX aruanne salvestatud: {out_xlsx}")

if __name__ == "__main__":
    main()