# ==============================================================
#  SOC Threat Analyser — v6.2 (2025-11-18)
#  Autor: ocrHeiki / GitHub (muudetud: assistent)
#  Kirjeldus:
#   - Analüüsib Palo Alto logisid (CSV)
#   - Toetab: 24h / 7d / 30d ajaaknaid
#   - Väljundid: DOCX, TXT, XLSX + graafikud (PNG)
#   - Funktsionaalsused:
#     * Severity, Threat, Action, IP (source/destination), Port, Trendid
#     * Threat + Port kaardistus
#     * DOCX raport koos visuaalidega ja sisukorraga
#     * Korrelatsioon: IP -> TOP Threats & TOP Kategooriad
# ==============================================================
from pathlib import Path
from datetime import datetime, timedelta
import pandas as pd
import matplotlib.pyplot as plt
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import json
import math

# ———————————————————————————————
# CLI seadistus
# ———————————————————————————————

def parse_args():
    parser = argparse.ArgumentParser(description="SOC Threat Analyser – Palo Alto logide multianalüüs")
    parser.add_argument("--timeframe", choices=["24h", "7d", "30d"], default="24h",
                        help="Analüüsitav ajaaken (vaikimisi: '24h')")
    return parser.parse_args()

# ———————————————————————————————
# Kaustad ja värvid
# ———————————————————————————————

BASE_DIR = Path.home() / "Documents" / "SOC"
RAW_DIR = BASE_DIR / "raw"
REPORTS_DIR = BASE_DIR / "reports"
RESULTS_DIR = BASE_DIR / "tulemused"
THREAT_VAULT_CACHE = BASE_DIR / "threat_vault_cache"
TRENDS_DIR = BASE_DIR / "trendid"

for d in [RAW_DIR, REPORTS_DIR, RESULTS_DIR, THREAT_VAULT_CACHE, TRENDS_DIR]:
    d.mkdir(parents=True, exist_ok=True)

COLORS_SEV = {"low": "#0000FF", "medium": "#FFFF00", "high": "#FFA500", "critical": "#FF0000"}
COLORS_ACTION = {"allow": "#33CC33", "deny": "#CC3333", "drop": "#3366CC", "alert": "#FFCC00", "reset-both": "#9933CC", "reset-server": "#800080"}
COLORS_CAT = {
    "command-and-control": "#CC0000", "code-execution": "#FF6600", "sql-injection": "#FF9933",
    "brute-force": "#FFCC00", "dos": "#FFFF66", "hacktool": "#9933CC", "info-leak": "#66CCFF",
    "spyware": "#3399FF", "code-obfuscation": "#996633"
}

# ———————————————————————————————
# Abifunktsioonid
# ———————————————————————————————

def iso_from_filename(name: str):
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", name)
    if m:
        return datetime.strptime(f"{m.group(3)}-{m.group(2)}-{m.group(1)}", "%Y-%m-%d").date()
    m = re.search(r"(\d{4})-(\d{2})-(\d{2})", name)
    return m.group(0) if m else datetime.now().date()

def first_existing(df, names):
    for n in names:
        if n in df.columns:
            return n
    return None

def norm_lower(col):
    return col.astype(str).str.lower().str.strip()

def overwrite_file(path: Path):
    if path.exists():
        try:
            path.unlink()
        except Exception as e:
            print(f"[~] Ei saanud kustutada faili: {path} ({e})")

def bar(series, title, outpath, colors=None, rot=0, is_date_trend=False):
    if series is None or series.empty:
        return
    plt.figure(figsize=(10, 5))
    try:
        cols = [colors.get(str(k).lower(), "#888888") for k in series.index] if colors else "#3399FF"
    except Exception:
        cols = "#3399FF"
    ax = series.plot(kind="bar", color=cols)
    plt.title(title)
    if is_date_trend:
        try:
            ax.set_xticklabels(series.index.strftime('%Y-%m-%d'))
        except:
            pass
        plt.xticks(rotation=45, ha='right')
    else:
        plt.xticks(rotation=rot)
    plt.tight_layout()
    overwrite_file(outpath)
    plt.savefig(outpath)
    plt.close()

def donut_with_side_legend(series, title, outpath, colors_map=None, legend_title="Legend"):
    if series is None or series.empty:
        return
    values = series.values
    labels = [str(i) for i in series.index]
    total = float(values.sum()) if values.sum() else 1.0
    pcts = [round(v * 100.0 / total, 1) for v in values]
    colors = [colors_map.get(str(i).lower(), "#888888") for i in series.index] if colors_map else None
    fig, ax = plt.subplots(figsize=(9, 6))
    wedges, _ = ax.pie(values, startangle=90, labels=None, colors=colors, wedgeprops=dict(width=0.38))
    ax.set_title(title)
    legend_lines = [f"{lbl} — {int(cnt)} ({pct}%)" for lbl, cnt, pct in zip(labels, values, pcts)]
    ax.legend(wedges, legend_lines, title=legend_title, loc="center left", bbox_to_anchor=(1.0, 0.5), frameon=False)
    plt.tight_layout()
    fig.savefig(outpath, bbox_inches="tight")
    plt.close(fig)

def add_image(doc, img_path, caption, width_in=6.0):
    if not img_path.exists():
        return
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ———————————————————————————————
# Threat Vault (cache-põhine)
# ———————————————————————————————

def get_threat_details(threat_name):
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', threat_name)
    cache_file = THREAT_VAULT_CACHE / f"{safe_name}.json"
    if cache_file.exists():
        try:
            with open(cache_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except: pass
    return {}

# ———————————————————————————————
# Threat + Portide seostamine
# ———————————————————————————————

def generate_threat_port_mapping(df_all):
    threat_col = first_existing(df_all, ["Threat/Content Name", "threat_name"])
    port_col = first_existing(df_all, ["Destination Port", "dest_port"])
    app_col = first_existing(df_all, ["Application", "app"])
    if not threat_col:
        print("[!] Threat Name veerg puudub.")
        return
    df_threats = df_all[df_all[threat_col].notna()].copy()

    def parse(text):
        m = re.match(r"^(.*?)\s*\((\d+)\)$", str(text).strip())
        return (m.group(1).strip(), m.group(2)) if m else (str(text).strip(), None)

    df_threats[["threat_clean", "threat_id"]] = df_threats[threat_col].apply(lambda x: pd.Series(parse(x)))

    mapping_file = RESULTS_DIR / "threat_port_mapping.txt"
    overwrite_file(mapping_file)

    with open(mapping_file, "w", encoding="utf-8") as f:
        f.write("# Palo Alto Threat / Port Mapping\n")
        f.write(f"# Genereeritud: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

        for name in df_threats["threat_clean"].dropna().unique():
            subset = df_threats[df_threats["threat_clean"] == name]
            threat_id = subset["threat_id"].iloc[0] if not subset["threat_id"].isna().all() else "N/A"

            f.write(f"Threat: \"{name} ({threat_id})\"\n")
            f.write(f"  Threat Vault ID: {threat_id}\n")

            if port_col and port_col in subset.columns:
                ports = subset[port_col].value_counts().head(10)
                f.write("  Portide kasutus:\n")
                for port, count in ports.items():
                    f.write(f"    {port} ({count} korda)\n")

            if app_col and app_col in subset.columns:
                apps = subset[app_col].value_counts().head(10)
                f.write("  Rakendused:\n")
                for app, count in apps.items():
                    f.write(f"    {app} ({count} korda)\n")
            f.write("-" * 50 + "\n\n")

    print(f"[+] Threat-port kaardistus salvestatud: {mapping_file}")

# ———————————————————————————————
# Trendianalüüs
# ———————————————————————————————

def analyze_trends(df_all):
    if "tname_norm" not in df_all.columns:
        print("[!] 'tname_norm' puudub – trendi ei saa teha.")
        return {}

    df = df_all.copy()
    df["log_date"] = pd.to_datetime(df["log_date"], errors='coerce')
    df = df.dropna(subset=["log_date"])

    df['week_start'] = df['log_date'].apply(lambda x: x - timedelta(days=x.weekday())).dt.normalize()
    weekly_volume = df.groupby('week_start').size()
    out_vol = TRENDS_DIR / "trend_weekly_volume.png"
    bar(weekly_volume, "Nädalapõhine Logide Maht", out_vol, is_date_trend=True)

    top5_names = df['tname_norm'].value_counts().head(5).index.tolist()
    weekly_threats = df.groupby(['week_start', 'tname_norm']).size().unstack(fill_value=0)
    top5_trend = weekly_threats[weekly_threats.columns.intersection(top5_names)]

    out_top5 = TRENDS_DIR / "trend_top5_threats.png"
    plt.figure(figsize=(12, 6))
    top5_trend.plot(kind='line', marker='o', ax=plt.gca())
    plt.title("TOP 5 Threat'i Aktiivsus Aja Jooksul")
    plt.xlabel("Nädala Algus")
    plt.ylabel("Kirjete Arv")
    plt.legend(title='Threat', bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.grid(True)
    plt.tight_layout()
    overwrite_file(out_top5)
    plt.savefig(out_top5)
    plt.close()

    print("[+] Trendianalüüs lõpetatud.")
    return {
        "plot_volume": out_vol,
        "plot_top5": out_top5
    }

# ———————————————————————————————
# IP -> Threat/Kategooria korrelatsiooni (abi)
# ———————————————————————————————

def ip_correlation(df, ip_col, threat_col, cat_col, top_n=10):
    """
    Tagastab DataFrame'i, kus iga IP-l on:
      - count (kui mitu kirjet)
      - top threats (näide: 'threat1(12), threat2(7)')
      - top categories (näide: 'cat1(8), cat2(3)')
    """
    if ip_col not in df.columns:
        return pd.DataFrame()

    counts = df[ip_col].value_counts().head(top_n)
    rows = []
    for ip, cnt in counts.items():
        sub = df[df[ip_col] == ip]
        top_th = sub[threat_col].value_counts().head(5) if threat_col and threat_col in sub.columns else pd.Series(dtype=int)
        top_cat = sub[cat_col].value_counts().head(5) if cat_col and cat_col in sub.columns else pd.Series(dtype=int)

        th_str = ", ".join([f"{str(t)}({int(c)})" for t, c in top_th.items()]) if not top_th.empty else ""
        cat_str = ", ".join([f"{str(ca)}({int(cc)})" for ca, cc in top_cat.items()]) if not top_cat.empty else ""

        rows.append({
            "ip": ip,
            "count": int(cnt),
            "top_threats": th_str,
            "top_categories": cat_str
        })
    return pd.DataFrame(rows)

# ———————————————————————————————
# Peamine logika
# ———————————————————————————————

def main():
    args = parse_args()
    timeframe_days = {"24h": 1, "7d": 7, "30d": 30}[args.timeframe]
    print(f"[i] Käivitan {args.timeframe} analüüsi...")

    csv_files = list(RAW_DIR.glob("*.csv"))
    if not csv_files:
        print("[!] Ühtegi CSV-faili ei leitud.")
        return

    dfs = []
    today = datetime.now().date()
    used_files = []

    for f in csv_files:
        try:
            df = pd.read_csv(f, low_memory=False)
        except Exception as e:
            print(f"[!] Viga faili lugemisel {f.name}: {e}")
            continue
        df["log_date"] = iso_from_filename(f.name)
        dfs.append(df)
        used_files.append(f.name)

    if not dfs:
        print("[!] Ühtegi loetavat CSV andmestikku pole.")
        return

    df_all = pd.concat(dfs, ignore_index=True)
    print(f"[i] Loetud {len(df_all)} kirjet {len(csv_files)} failist.")

    # --- veerud (leiame parimad võimalikud nimetused)
    threat_col = first_existing(df_all, ["Threat/Content Name", "threat_name"])
    sev_col    = first_existing(df_all, ["Severity", "severity"])
    act_col    = first_existing(df_all, ["Action", "action"])
    src_col    = first_existing(df_all, ["Source address", "src", "source"])
    dst_col    = first_existing(df_all, ["Destination address", "dst", "destination"])
    cat_col    = first_existing(df_all, ["thr_category", "category"])
    type_col   = first_existing(df_all, ["Threat/Content Type", "threat_type"])

    # normaliseerime threat nime (tname_norm) kui on
    if threat_col:
        df_all["tname_norm"] = norm_lower(df_all[threat_col])
    else:
        df_all["tname_norm"] = "unknown"

    # normaliseeri muud veerud
    df_all["sev_norm"] = norm_lower(df_all[sev_col]) if sev_col else pd.Series(["unknown"] * len(df_all))
    df_all["act_norm"] = norm_lower(df_all[act_col]) if act_col else pd.Series(["unknown"] * len(df_all))
    df_all["src_norm"] = df_all[src_col].astype(str).str.strip() if src_col else pd.Series(["unknown"] * len(df_all))
    df_all["dst_norm"] = df_all[dst_col].astype(str).str.strip() if dst_col else pd.Series(["unknown"] * len(df_all))
    if cat_col and cat_col in df_all.columns:
        df_all["cat_norm"] = norm_lower(df_all[cat_col])
    else:
        df_all["cat_norm"] = pd.Series(["unknown"] * len(df_all))

    # filter aeg
    df_all["log_date"] = pd.to_datetime(df_all["log_date"], errors='coerce')
    cutoff = today - timedelta(days=timeframe_days)
    df_filtered = df_all[df_all["log_date"].dt.date >= cutoff]

    if df_filtered.empty:
        print("[!] Filtritud andmed on tühjad.")
        return

    # kokkuvõtted
    sev_counts = df_filtered["sev_norm"].value_counts()
    action_counts = df_filtered["act_norm"].value_counts()
    top_threats = df_filtered["tname_norm"].value_counts().head(10)
    top_src_ips = df_filtered["src_norm"].value_counts().head(10)
    top_dst_ips = df_filtered["dst_norm"].value_counts().head(10)
    cat_counts = df_filtered["cat_norm"].value_counts().head(10)

    # graafikufailid
    plots = {
        "severity_donut": REPORTS_DIR / "severity_donut.png",
        "severity_bar": REPORTS_DIR / "severity_bar.png",
        "action_bar": REPORTS_DIR / "action_bar.png",
        "action_pie": REPORTS_DIR / "action_pie.png",
        "top_threats": REPORTS_DIR / "top_threats.png",
        "top_src_ips": REPORTS_DIR / "top_src_ips.png",
        "top_dst_ips": REPORTS_DIR / "top_dst_ips.png",
        "top_categories": REPORTS_DIR / "top_categories.png",
    }

    # genereeri graafikud
    donut_with_side_legend(sev_counts, "Severity osakaal", plots["severity_donut"], COLORS_SEV, "Sev – Arv (osakaal)")
    bar(sev_counts, "Severity Jaotus", plots["severity_bar"], COLORS_SEV)
    bar(action_counts, "Action Jaotus", plots["action_bar"], COLORS_ACTION, rot=45)
    if not action_counts.empty:
        plt.figure(figsize=(6,6))
        plt.pie(action_counts, labels=action_counts.index, autopct="%1.1f%%", startangle=90)
        plt.title("Action osakaal")
        overwrite_file(plots["action_pie"])
        plt.savefig(plots["action_pie"])
        plt.close()
    bar(top_threats, "Top Threats", plots["top_threats"], rot=45)
    bar(top_src_ips, "Top Source IPs", plots["top_src_ips"], rot=45)
    bar(top_dst_ips, "Top Destination IPs", plots["top_dst_ips"], rot=45)
    if not cat_counts.empty:
        bar(cat_counts, "Top Kategooriad", plots["top_categories"], COLORS_CAT, rot=45)

    # threat-port kaartistus + trendid
    generate_threat_port_mapping(df_all)
    trend_data = analyze_trends(df_all) if timeframe_days > 1 else {}

    # ---------------------------
    # IP korrelatsioonid (Source & Destination)
    # ---------------------------
    src_corr_df = ip_correlation(df_filtered, "src_norm", "tname_norm", "cat_norm", top_n=10) if "src_norm" in df_filtered.columns else pd.DataFrame()
    dst_corr_df = ip_correlation(df_filtered, "dst_norm", "tname_norm", "cat_norm", top_n=10) if "dst_norm" in df_filtered.columns else pd.DataFrame()

    # ---------------------------
    # TXT raport
    # ---------------------------
    txt_path = RESULTS_DIR / f"soc_{args.timeframe}_summary.txt"
    overwrite_file(txt_path)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(f"SOC {args.timeframe.upper()} ARUANNE – {today.strftime('%d.%m.%Y')}\n")
        f.write("=" * 60 + "\n")
        f.write(f"- Kirjeid (filtritud): {len(df_filtered)}\n")
        if not top_threats.empty:
            f.write(f"- Populaarseim threat: {top_threats.index[0]} ({top_threats.iloc[0]}x)\n")
        if not top_src_ips.empty:
            f.write(f"- Populaarseim Source IP: {top_src_ips.index[0]} ({top_src_ips.iloc[0]}x)\n")
        if not top_dst_ips.empty:
            f.write(f"- Populaarseim Destination IP: {top_dst_ips.index[0]} ({top_dst_ips.iloc[0]}x)\n")
        f.write(f"- Kasutatud CSV fail(id): {', '.join(used_files)}\n\n")

        # Severity
        f.write("■ Severity jaotus:\n")
        if not sev_counts.empty:
            sev_order = ["low", "medium", "high", "critical"]
            for s in sev_order:
                if s in sev_counts.index:
                    f.write(f"  - {s.capitalize():<9}: {sev_counts[s]}\n")
            for s, cnt in sev_counts.items():
                if s not in sev_order:
                    f.write(f"  - {str(s).capitalize():<9}: {cnt}\n")
        else:
            f.write("  - Puuduvad andmed\n")

        # Action
        f.write("\n■ Action jaotus:\n")
        if not action_counts.empty:
            for act, cnt in action_counts.items():
                f.write(f"  - {act:<12}: {cnt}\n")
        else:
            f.write("  - Puuduvad andmed\n")

        # TOP kategooriad
        f.write("\n■ TOP kategooriad:\n")
        if not cat_counts.empty:
            for i, (cat, cnt) in enumerate(cat_counts.items(), 1):
                f.write(f"  {i}. {cat} – {cnt}\n")
        else:
            f.write("  - Puuduvad andmed\n")

        # TOP threats
        f.write("\nTOP 10 Threati:\n")
        for i, (t, c) in enumerate(top_threats.items(), 1):
            f.write(f"  {i}. {t} – {c}x\n")

        # TOP source IPs + korrelatsioon
        f.write("\nTOP 10 SOURCE IP aadressid (algpunktid):\n")
        if not src_corr_df.empty:
            for _, row in src_corr_df.iterrows():
                f.write(f"  - {row['ip']} ({row['count']}x)\n")
                if row['top_threats']:
                    f.write(f"      TOP Threats: {row['top_threats']}\n")
                if row['top_categories']:
                    f.write(f"      TOP Categories: {row['top_categories']}\n")
        else:
            f.write("  - Puuduvad andmed\n")

        # TOP destination IPs + korrelatsioon
        f.write("\nTOP 10 DESTINATION IP aadressid (sihtmärgid):\n")
        if not dst_corr_df.empty:
            for _, row in dst_corr_df.iterrows():
                f.write(f"  - {row['ip']} ({row['count']}x)\n")
                if row['top_threats']:
                    f.write(f"      TOP Threats: {row['top_threats']}\n")
                if row['top_categories']:
                    f.write(f"      TOP Categories: {row['top_categories']}\n")
        else:
            f.write("  - Puuduvad andmed\n")

    # ---------------------------
    # XLSX raport (sheet per tabel)
    # ---------------------------
    xlsx_path = RESULTS_DIR / f"soc_{args.timeframe}_report.xlsx"
    overwrite_file(xlsx_path)
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
        # top threats
        top_threats.rename_axis("threat").reset_index(name="count").to_excel(writer, sheet_name="Top_Threats", index=False)
        # severity
        sev_counts.rename_axis("severity").reset_index(name="count").to_excel(writer, sheet_name="Severity", index=False)
        # action
        action_counts.rename_axis("action").reset_index(name="count").to_excel(writer, sheet_name="Action", index=False)
        # top categories
        cat_counts.rename_axis("category").reset_index(name="count").to_excel(writer, sheet_name="Top_Categories", index=False)
        # top src/dst simple
        top_src_ips.rename_axis("src_ip").reset_index(name="count").to_excel(writer, sheet_name="Top_Source_IPs", index=False)
        top_dst_ips.rename_axis("dst_ip").reset_index(name="count").to_excel(writer, sheet_name="Top_Destination_IPs", index=False)
        # detailed correlations
        if not src_corr_df.empty:
            src_corr_df.to_excel(writer, sheet_name="Src_IP_Correlation", index=False)
        if not dst_corr_df.empty:
            dst_corr_df.to_excel(writer, sheet_name="Dst_IP_Correlation", index=False)
    print(f"[+] XLSX salvestatud: {xlsx_path}")

    # ---------------------------
    # DOCX raport
    # ---------------------------
    docx_path = RESULTS_DIR / f"soc_{args.timeframe}_report.docx"
    overwrite_file(docx_path)
    doc = Document()
    doc.add_heading(f"SOC {args.timeframe.upper()} ARUANNE – {today.strftime('%d.%m.%Y')}", 0)

    # Sisukord
    doc.add_heading("Sisukord", level=1)
    doc.add_paragraph("1. Lühikokkuvõte", style="List Bullet")
    doc.add_paragraph("2. Detailne Ülevaade", style="List Bullet")
    doc.add_paragraph("3. TOP IP Korrelatsioonid (Source & Destination)", style="List Bullet")
    doc.add_paragraph("4. Graafikud", style="List Bullet")
    doc.add_paragraph("5. Lisateave: Threat + Port", style="List Bullet")

    # Lühikokkuvõte
    doc.add_heading("1. Lühikokkuvõte", level=1)
    doc.add_paragraph(f"📈 Kirjete arv (filtritud): {len(df_filtered)}")
    if not top_threats.empty:
        doc.add_paragraph(f"🔥 Aktiivseim threat: {top_threats.index[0]} ({top_threats.iloc[0]}x)")
    if not top_src_ips.empty:
        doc.add_paragraph(f"📡 Populaarseim Source IP: {top_src_ips.index[0]} ({top_src_ips.iloc[0]}x)")
    if not top_dst_ips.empty:
        doc.add_paragraph(f"🎯 Populaarseim Destination IP: {top_dst_ips.index[0]} ({top_dst_ips.iloc[0]}x)")
    doc.add_paragraph(f"📁 Kasutatud failid: {', '.join(used_files)}")

    # Detailne ülevaade
    doc.add_heading("2. Detailne Ülevaade", level=1)
    doc.add_paragraph(f"- Kirjeid: {len(df_filtered)}")
    if not top_threats.empty:
        doc.add_paragraph(f"- Populaarseim threat: {top_threats.index[0]} ({top_threats.iloc[0]}x)")
    # Severity
    doc.add_paragraph("\n■ Severity jaotus:")
    if not sev_counts.empty:
        sev_order = ["low", "medium", "high", "critical"]
        for s in sev_order:
            if s in sev_counts.index:
                doc.add_paragraph(f"  - {s.capitalize():<9}: {sev_counts[s]}")
        for s, cnt in sev_counts.items():
            if s not in sev_order:
                doc.add_paragraph(f"  - {str(s).capitalize():<9}: {cnt}")
    else:
        doc.add_paragraph("  - Puuduvad andmed")
    # Action
    doc.add_paragraph("\n■ Action jaotus:")
    if not action_counts.empty:
        for act, cnt in action_counts.items():
            doc.add_paragraph(f"  - {act:<12}: {cnt}")
    else:
        doc.add_paragraph("  - Puuduvad andmed")
    # Top categories
    doc.add_paragraph("\n■ TOP kategooriad:")
    if not cat_counts.empty:
        for i, (cat, cnt) in enumerate(cat_counts.items(), 1):
            doc.add_paragraph(f"  {i}. {cat} – {cnt}")
    else:
        doc.add_paragraph("  - Puuduvad andmed")

    # IP korrelatsioonid section
    doc.add_heading("3. TOP IP Korrelatsioonid", level=1)
    doc.add_paragraph("Allpool on eraldi loendid Source (algpunktid) ja Destination (sihtmärgid). Iga IP juures on näidatud TOP Threats ja TOP Kategooriad selle IP-ga seotud kirjetest.")

    # Source IPs
    doc.add_paragraph("\nTOP 10 SOURCE IP aadressid (algpunktid):")
    if not src_corr_df.empty:
        for _, row in src_corr_df.iterrows():
            doc.add_paragraph(f"  - {row['ip']} ({row['count']}x)")
            if row['top_threats']:
                doc.add_paragraph(f"      TOP Threats: {row['top_threats']}")
            if row['top_categories']:
                doc.add_paragraph(f"      TOP Categories: {row['top_categories']}")
    else:
        doc.add_paragraph("  - Puuduvad andmed")

    # Destination IPs
    doc.add_paragraph("\nTOP 10 DESTINATION IP aadressid (sihtmärgid):")
    if not dst_corr_df.empty:
        for _, row in dst_corr_df.iterrows():
            doc.add_paragraph(f"  - {row['ip']} ({row['count']}x)")
            if row['top_threats']:
                doc.add_paragraph(f"      TOP Threats: {row['top_threats']}")
            if row['top_categories']:
                doc.add_paragraph(f"      TOP Categories: {row['top_categories']}")
    else:
        doc.add_paragraph("  - Puuduvad andmed")

    # Graafikud
    doc.add_heading("4. Graafikud", level=1)
    add_image(doc, plots["severity_donut"], "Severity osakaal (sõõrik)")
    add_image(doc, plots["severity_bar"], "Severity jaotus")
    add_image(doc, plots["action_bar"], "Action jaotus")
    add_image(doc, plots["action_pie"], "Action jaotus (pie)")
    add_image(doc, plots["top_threats"], "Top Threats")
    add_image(doc, plots["top_src_ips"], "Top Source IP-d")
    add_image(doc, plots["top_dst_ips"], "Top Destination IP-d")
    add_image(doc, plots["top_categories"], "Top Kategooriad")

    # Trendid (kui tehtud)
    if trend_data:
        doc.add_heading("5. Trendid", level=1)
        add_image(doc, trend_data["plot_volume"], "Nädalane logide maht")
        add_image(doc, trend_data["plot_top5"], "TOP 5 threati aktiivsus ajas")

    # Threat-port kaardistus
    port_map_path = RESULTS_DIR / "threat_port_mapping.txt"
    if port_map_path.exists():
        doc.add_heading("6. Lisateave: Threat + Port Kaardistus", level=1)
        try:
            with open(port_map_path, "r", encoding="utf-8") as f:
                content = f.read()
            doc.add_paragraph(content[:8000] + ("..." if len(content) > 8000 else ""))
        except Exception as e:
            doc.add_paragraph(f"❌ Viga: {e}")

    doc.save(docx_path)
    print(f"[OK] DOCX raport salvestatud: {docx_path}")
    print(f"[OK] TXT raport salvestatud: {txt_path}")

if __name__ == "__main__":
    main()
