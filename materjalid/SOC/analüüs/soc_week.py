# ==============================================================
#  SOC 7 päeva koondanalüüs — v3.0
#  Autor: ocrHeiki (GitHub: https://github.com/ocrHeiki)
#  Kirjeldus:
#   - Leiab ja kasutab ainult neid CSV-faile, mille sisu katab ~5–7 päeva
#   - Välistab automaatselt 24h failid (1–2 päeva katvusega)
#   - Loob TXT, XLSX, DOCX aruanded + PNG graafikud (legendidega)
#   - Lisab TOP 10 Threat/Content Name (tekst + graafik)
#   - Kirjutab week_ip_YYYY-MM-DD.txt (TOP 10 lähte-IP)
# ==============================================================

from pathlib import Path
from datetime import datetime
import pandas as pd
import matplotlib
matplotlib.use("Agg")  # headless
import matplotlib.pyplot as plt
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError

# --- Kaustad ---
BASE = Path.home() / "Documents" / "SOC"
RAW = BASE / "raw"
OUT = BASE / "tulemused"
REP = BASE / "reports"
for d in (RAW, OUT, REP):
    d.mkdir(parents=True, exist_ok=True)

# --- Värvikaardid ---
COLORS_SEV = {"low":"#0000FF","medium":"#FFFF00","high":"#FFA500","critical":"#FF0000"}
COLORS_ACTION = {"allow":"#33CC33","deny":"#CC3333","drop":"#3366CC","alert":"#FFCC00","reset-both":"#9933CC","reset-server":"#800080"}
COLORS_TYPE = {"malware":"#CC0033","vulnerability":"#FF3333","spyware":"#3399FF","suspicious":"#FFCC66","benign":"#66CC66"}
COLORS_CAT = {
    "command-and-control":"#CC0000","code-execution":"#FF6600","sql-injection":"#FF9933","brute-force":"#FFCC00",
    "dos":"#FFFF66","hacktool":"#9933CC","info-leak":"#66CCFF","spyware":"#3399FF","code-obfuscation":"#996633"
}

# --- Abid ---
def iso_from_filename(name: str) -> str:
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", name)  # dd.mm.yyyy
    if m: return f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
    m2 = re.search(r"(\d{4})-(\d{2})-(\d{2})", name)   # yyyy-mm-dd
    return m2.group(0) if m2 else datetime.now().date().isoformat()

def first_existing(df: pd.DataFrame, names) -> str | None:
    for n in names:
        if n in df.columns: return n
    return None

def norm_lower(s: pd.Series) -> pd.Series:
    return s.astype(str).str.strip().str.lower()

def safe_colors(series: pd.Series, cmap: dict[str,str] | None, fallback="#888888"):
    if cmap is None:
        return fallback
    return [cmap.get(str(i).lower(), fallback) for i in series.index]

def bar(series, title, outpath: Path, colors=None, rot=0, ylabel=None, legend_label=None):
    if series is None or series.empty: return
    plt.figure(figsize=(10,5))
    c = safe_colors(series, colors)
    series.plot(kind="bar", color=c, label=legend_label)
    plt.title(title)
    if ylabel: plt.ylabel(ylabel)
    if legend_label: plt.legend()
    plt.xticks(rotation=rot, ha="right" if rot else "center")
    plt.tight_layout()
    outpath.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(outpath)
    plt.close()

def pie(series, title, outpath: Path):
    if series is None or series.empty: return
    plt.figure(figsize=(7,7))
    # kasutame plt.pie, et sildid oleks alati loetavad (mitte pandas.plot(kind="pie"))
    wedges, texts, autotexts = plt.pie(
        series.values,
        labels=[str(i) for i in series.index],
        autopct="%1.1f%%",
        startangle=90,
        pctdistance=0.8,
        labeldistance=1.05
    )
    plt.title(title)
    plt.tight_layout()
    outpath.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(outpath)
    plt.close()

def add_image_safe(doc: Document, img_path: Path, caption: str = "", width_in=6.0):
    try:
        if img_path.exists() and img_path.is_file():
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(width_in))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"(Märkus: graafik puudub) {img_path.name}")
    except (PackageNotFoundError, OSError) as e:
        doc.add_paragraph(f"(Märkus: ei saanud lisada {img_path.name} – {e})")

def parse_time_columns(df: pd.DataFrame) -> pd.Series | None:
    """Püüa leida ajatempli veerg ja tagasta kuupäev (YYYY-MM-DD) seeriat."""
    time_cols = [
        "Time Generated","@timestamp","Receive Time","Event Received Time",
        "time","timestamp","Event Time","Generate Time","log_time"
    ]
    for col in time_cols:
        if col in df.columns:
            ts = pd.to_datetime(df[col], errors="coerce", utc=True).dropna()
            if not ts.empty:
                return ts.dt.tz_convert(None).dt.date  # naivseks päevaks
    return None

def file_span_dates(df: pd.DataFrame, fallback_iso: str) -> tuple[str,str,int]:
    """Tagasta (start_date, end_date, days_covered). Kui timestamp puudub, 1-päevane katvus fallback_iso põhjal."""
    dates = parse_time_columns(df)
    if dates is not None and len(dates) > 0:
        s = str(min(dates))
        e = str(max(dates))
        days = (pd.to_datetime(e) - pd.to_datetime(s)).days + 1
        return s, e, days
    # fallback: ainult failinime kuupäev
    return fallback_iso, fallback_iso, 1

def load_day(path: Path) -> tuple[pd.DataFrame, dict]:
    """Laadi CSV, normaliseeri veerud, tagasta df + meta (period)."""
    try:
        df = pd.read_csv(path, encoding="utf-8", low_memory=False)
    except UnicodeDecodeError:
        df = pd.read_csv(path, encoding="latin-1", low_memory=False)

    sev_col  = first_existing(df, ["Severity","severity"])
    risk_col = first_existing(df, ["Risk","risk","Risk Level","risk_level","Risk of app"])
    cat_col  = first_existing(df, ["thr_category","category","Threat Category"])
    type_col = first_existing(df, ["Threat/Content Type","Threat Type","threat_type","threat/content type"])
    act_col  = first_existing(df, ["Action","action"])
    name_col = first_existing(df, ["Threat/Content Name","Threat Name","content_name","threat_name"])
    src_col  = first_existing(df, ["Source address","Src","Source","src_ip","Source IP","src"])
    dst_col  = first_existing(df, ["Destination address","Dst","Destination","dst_ip","Destination IP","dst"])

    if sev_col:  df["sev_norm"]   = norm_lower(df[sev_col])
    if risk_col: df["risk_norm"]  = pd.to_numeric(df[risk_col].astype(str).str.extract(r"(\d+)")[0], errors="coerce").fillna(0).astype(int)
    if cat_col:  df["cat_norm"]   = norm_lower(df[cat_col])
    if type_col: df["type_norm"]  = norm_lower(df[type_col])
    if act_col:  df["act_norm"]   = norm_lower(df[act_col])
    if name_col: df["tname_norm"] = df[name_col].astype(str).str.strip()
    if src_col:  df["src_norm"]   = df[src_col].astype(str).str.strip()
    if dst_col:  df["dst_norm"]   = df[dst_col].astype(str).str.strip()

    # Perioodi leidmine failist
    fallback = iso_from_filename(path.name)
    start, end, days = file_span_dates(df, fallback)
    meta = {"start": start, "end": end, "days": days}
    return df, meta

# --- Peafunktsioon ---
def main():
    csv_files = sorted(RAW.glob("*.csv"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not csv_files:
        print("[!] Ei leitud CSV-faile kaustas 'raw/'.")
        return

    # Lae failid ja arvuta katvus
    loaded = []
    for p in csv_files:
        df, meta = load_day(p)
        loaded.append((p, df, meta))

    # Filtreeri välja "nädalafailid": katvus 5..7 (lubame kuni 9, kui piiripealne)
    week_like = [(p,df,m) for (p,df,m) in loaded if 5 <= m["days"] <= 9]
    # Välista 24h failid (<=2 päeva)
    week_like = [(p,df,m) for (p,df,m) in week_like if m["days"] >= 5]

    if not week_like:
        # Fallback: võta vähemalt 2 viimast suuremat faili, et üldse midagi oleks
        print("[WARN] Nädalakatvusega faile ei leitud. Kasutan viimaseid suuremaid faile.")
        # sorteeri katvuse järgi
        tmp = sorted(loaded, key=lambda x: x[2]["days"], reverse=True)[:2]
        week_like = tmp

    # Sorteeri failid perioodi alguse järgi (kronoloogiline)
    week_like = sorted(week_like, key=lambda x: x[2]["start"])

    # Koonda kõik read ühte nädalaraami
    all_df = pd.concat([df for _, df, _ in week_like], ignore_index=True) if week_like else pd.DataFrame()

    # Päevade loetav vahemik aruande päisesse
    period_start = week_like[0][2]["start"]
    period_end   = week_like[-1][2]["end"]

    # Päevane jaotus (kui ajatempli veerg olemas)
    dates_series = parse_time_columns(all_df)
    if dates_series is None:
        # Fallback: kasuta iga faili start-kuupäeva; pane ridu prooviks "kuupäev" tulpa
        all_df["date"] = period_start
    else:
        all_df["date"] = dates_series.astype(str)

    # Koonda nädalastatistika
    sev_counts = all_df["sev_norm"].value_counts() if "sev_norm" in all_df.columns else pd.Series(dtype=int)
    type_counts = all_df["type_norm"].value_counts() if "type_norm" in all_df.columns else pd.Series(dtype=int)
    act_counts = all_df["act_norm"].value_counts() if "act_norm" in all_df.columns else pd.Series(dtype=int)
    cat_counts = all_df["cat_norm"].value_counts() if "cat_norm" in all_df.columns else pd.Series(dtype=int)
    top_src = all_df["src_norm"].value_counts().head(10) if "src_norm" in all_df.columns else pd.Series(dtype=int)
    top_dst = all_df["dst_norm"].value_counts().head(10) if "dst_norm" in all_df.columns else pd.Series(dtype=int)
    top_threats = all_df["tname_norm"].value_counts().head(10) if "tname_norm" in all_df.columns else pd.Series(dtype=int)

    # Päeva x severity "stacked" ettevalmistus
    sev_order = ["low","medium","high","critical"]
    if "sev_norm" in all_df.columns and "date" in all_df.columns:
        day_sev = (all_df.groupby(["date","sev_norm"]).size()
                   .unstack(fill_value=0)
                   .reindex(columns=sev_order, fill_value=0)
                   .sort_index())
    else:
        day_sev = pd.DataFrame()

    # Päevapõhine kokkuvõte
    day_summary = (
        all_df.groupby("date").size().rename("alerts_total").to_frame()
        if "date" in all_df.columns else pd.DataFrame()
    )
    if not day_summary.empty and "sev_norm" in all_df.columns:
        for s in sev_order:
            day_summary[s] = all_df[all_df["sev_norm"]==s].groupby("date").size()
        day_summary = day_summary.fillna(0).astype(int)
        day_summary["hi_crit"] = day_summary.get("high",0) + day_summary.get("critical",0)
        day_summary["hi_crit_pct"] = (day_summary["hi_crit"] / day_summary["alerts_total"] * 100).round(2)

    # Võrdlus esimene pool vs teine pool (vajab >=2 faili)
    def top5_diff(curr: pd.Series, prev: pd.Series):
        idx = set(curr.index.astype(str)).union(set(prev.index.astype(str)))
        c = curr.reindex(idx, fill_value=0).astype(int)
        p = prev.reindex(idx, fill_value=0).astype(int)
        diff = (c - p).sort_values(ascending=False)
        return diff.head(5), diff.tail(5)

    if len(week_like) >= 2:
        mid = len(week_like)//2
        first_df = pd.concat([df for _,df,_ in week_like[:mid]], ignore_index=True)
        last_df  = pd.concat([df for _,df,_ in week_like[mid:]], ignore_index=True)
        prev_cat = first_df["cat_norm"].value_counts() if "cat_norm" in first_df.columns else pd.Series(dtype=int)
        prev_src = first_df["src_norm"].value_counts() if "src_norm" in first_df.columns else pd.Series(dtype=int)
        prev_dst = first_df["dst_norm"].value_counts() if "dst_norm" in first_df.columns else pd.Series(dtype=int)
        curr_cat = last_df["cat_norm"].value_counts() if "cat_norm" in last_df.columns else pd.Series(dtype=int)
        curr_src = last_df["src_norm"].value_counts() if "src_norm" in last_df.columns else pd.Series(dtype=int)
        curr_dst = last_df["dst_norm"].value_counts() if "dst_norm" in last_df.columns else pd.Series(dtype=int)
        cat_up, cat_down = top5_diff(curr_cat, prev_cat)
        src_up, src_down = top5_diff(curr_src, prev_src)
        dst_up, dst_down = top5_diff(curr_dst, prev_dst)
    else:
        cat_up = cat_down = src_up = src_down = dst_up = dst_down = pd.Series(dtype=int)

    # --- Väljundfailide nimed ---
    today = datetime.now().date().isoformat()
    out_txt = OUT / f"week_summary_{today}.txt"
    out_xlsx = OUT / f"week_summary_{today}.xlsx"
    ip_txt = OUT / f"week_ip_{today}.txt"

    # --- TXT ---
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write(f"SOC NÄDALA KOONDARUANNE — {period_start} … {period_end}\n")
        f.write("="*72 + "\n\n")
        f.write(f"Kasutatud faile: {len(week_like)} tk\n")
        f.write("Allikafailid (kronoloogiliselt, vanem → uuem):\n")
        for p, _, meta in week_like:
            per = meta
            if per["start"] == per["end"]:
                f.write(f" - {per['start']}: {p.name}\n")
            else:
                f.write(f" - {per['start']} … {per['end']}: {p.name}\n")

        f.write("\n● Päevade ülevaade (alerts kokku / High+Critical / %):\n")
        if not day_summary.empty:
            f.write(f"{'Kuupäev':<12} {'Kokku':>7} {'High+Crit':>10} {'%':>6}\n")
            f.write("-"*42 + "\n")
            for d, r in day_summary.sort_index().iterrows():
                f.write(f"{d:<12} {int(r['alerts_total']):>7} {int(r.get('hi_crit',0)):>10} {float(r.get('hi_crit_pct',0)):>6.1f}\n")
        else:
            f.write(" (Puudub ajatempli info; päevade lõikes jaotust ei saanud arvutada)\n")

        if not cat_counts.empty:
            f.write("\n● TOP kategooriad (nädal):\n")
            for i, (k, v) in enumerate(cat_counts.head(10).items(), 1):
                f.write(f"  {i}. {k} — {int(v)}\n")

        if not top_threats.empty:
            f.write("\n● TOP 10 Threat / Content Name:\n")
            for i, (k, v) in enumerate(top_threats.items(), 1):
                f.write(f"  {i}. {k} — {int(v)}\n")

        if not top_src.empty:
            f.write("\n● TOP 10 allika IP:\n")
            for i,(k,v) in enumerate(top_src.items(),1):
                f.write(f"  {i}. {k} — {int(v)}\n")

        if not top_dst.empty:
            f.write("\n● TOP 10 sihtmärgi IP:\n")
            for i,(k,v) in enumerate(top_dst.items(),1):
                f.write(f"  {i}. {k} — {int(v)}\n")

        # Võrdlusplokk
        def write_diff(title, up, down):
            f.write(f"\n● {title} — võrdlus (esimene pool vs teine pool):\n")
            empty = True
            if up is not None and not up.empty:
                empty = False
                f.write("  ↑ Tõusud:\n")
                for i,(k,v) in enumerate(up.items(),1):
                    f.write(f"    {i}. {k}  (+{int(v)})\n")
            if down is not None and not down.empty:
                empty = False
                f.write("  ↓ Langused:\n")
                for i,(k,v) in enumerate(down.items(),1):
                    f.write(f"    {i}. {k}  ({int(v)})\n")
            if empty:
                f.write("  – (Võrdlemiseks vajatakse vähemalt 2 nädalafaili)\n")

        write_diff("Kategooriad", cat_up, cat_down)
        write_diff("Allika IP",  src_up, src_down)
        write_diff("Sihtmärgi IP", dst_up, dst_down)

    # --- IP nimekiri (TOP 10 lähte-IP) ---
    with open(ip_txt, "w", encoding="utf-8") as f:
        f.write("TOP 10 Allika IP (nädal)\n")
        f.write("="*30 + "\n\n")
        if not top_src.empty:
            for ip, cnt in top_src.items():
                f.write(f"{ip} ({cnt} korda)\n")
        else:
            f.write("(Andmeid ei leitud)\n")

    # --- XLSX ---
    with pd.ExcelWriter(out_xlsx) as xw:
        info_df = pd.DataFrame({
            "Key": ["Periood algus","Periood lõpp","Faile","Genereeritud"],
            "Value": [period_start, period_end, len(week_like), datetime.now().strftime("%Y-%m-%d %H:%M:%S")]
        })
        info_df.to_excel(xw, sheet_name="Info", index=False)

        if not day_summary.empty:
            day_summary.reset_index().rename(columns={"index":"date"}).to_excel(xw, sheet_name="DaySummary", index=False)
        if not sev_counts.empty:
            sev_counts.rename_axis("severity").reset_index(name="count").to_excel(xw, sheet_name="Severity", index=False)
        if not type_counts.empty:
            type_counts.rename_axis("type").reset_index(name="count").to_excel(xw, sheet_name="ThreatType", index=False)
        if not act_counts.empty:
            act_counts.rename_axis("action").reset_index(name="count").to_excel(xw, sheet_name="Action", index=False)
        if not cat_counts.empty:
            cat_counts.rename_axis("category").reset_index(name="count").to_excel(xw, sheet_name="Categories", index=False)
        if not top_threats.empty:
            top_threats.rename_axis("threat_name").reset_index(name="count").to_excel(xw, sheet_name="TopThreats", index=False)
        if not top_src.empty:
            top_src.rename_axis("src_ip").reset_index(name="count").to_excel(xw, sheet_name="TopSrc", index=False)
        if not top_dst.empty:
            top_dst.rename_axis("dst_ip").reset_index(name="count").to_excel(xw, sheet_name="TopDst", index=False)

    # --- Graafikud ---
    # Trend (päevad)
    trend_png = REP / f"week_trend_{today}.png"
    if not day_summary.empty:
        plt.figure(figsize=(11,6))
        plt.plot(day_summary.index, day_summary["alerts_total"], marker="o", label="Alerts kokku")
        if "hi_crit" in day_summary.columns:
            plt.plot(day_summary.index, day_summary["hi_crit"], marker="o", label="High+Critical")
        if "hi_crit_pct" in day_summary.columns:
            ax = plt.gca()
            ax2 = ax.twinx()
            ax2.plot(day_summary.index, day_summary["hi_crit_pct"], marker="o", linestyle="--", label="% High+Critical", color="#CC0000")
            ax2.set_ylabel("% High+Critical")
            # kahekordse telje legend:
            lines1, labels1 = ax.get_legend_handles_labels()
            lines2, labels2 = ax2.get_legend_handles_labels()
            plt.legend(lines1+lines2, labels1+labels2, loc="upper left")
        else:
            plt.legend()
        plt.title("■ Nädala trend: Alerts vs High+Critical")
        plt.xticks(rotation=45, ha="right")
        plt.tight_layout()
        plt.savefig(trend_png)
        plt.close()

    # Stacked severity päevade lõikes
    sev_png = REP / f"week_severity_stacked_{today}.png"
    if not day_sev.empty:
        plt.figure(figsize=(12,6))
        bottom = None
        for sev in sev_order:
            vals = day_sev[sev] if sev in day_sev.columns else 0
            plt.bar(day_sev.index, vals, bottom=bottom, label=sev.title(), color=COLORS_SEV.get(sev,"#888888"))
            bottom = (bottom + vals) if bottom is not None else vals
        plt.title("■ Severity jaotus päevade lõikes (stacked)")
        plt.xticks(rotation=45, ha="right")
        plt.legend()
        plt.tight_layout()
        plt.savefig(sev_png)
        plt.close()

    # Muud graafikud
    if not cat_counts.empty:
        bar(cat_counts.head(10), f"■ TOP 10 kategooriat (nädal) – {today}", REP / f"week_top_categories_{today}.png", colors=COLORS_CAT, rot=45)

    if not act_counts.empty:
        bar(act_counts, f"■ Action jaotus (nädal) – {today}", REP / f"week_action_bar_{today}.png", colors=COLORS_ACTION, rot=45, legend_label="Action", ylabel="Count")
        pie(act_counts, f"▲ Action osakaal (nädal) – {today}", REP / f"week_action_pie_{today}.png")

    if not type_counts.empty:
        pie(type_counts, f"▲ Threat/Content Type (nädal) – {today}", REP / f"week_threat_type_pie_{today}.png")

    if not top_src.empty:
        bar(top_src, f"■ TOP 10 allika IP (nädal) – {today}", REP / f"week_top_src_{today}.png", rot=45, ylabel="Count")

    if not top_dst.empty:
        bar(top_dst, f"■ TOP 10 sihtmärgi IP (nädal) – {today}", REP / f"week_top_dst_{today}.png", rot=45, ylabel="Count")

    if not top_threats.empty:
        bar(top_threats, f"■ TOP 10 Threat / Content Name (nädal) – {today}", REP / f"week_top_threats_{today}.png", rot=45, ylabel="Count")

    # --- DOCX ---
    docx_path = OUT / f"week_summary_{today}.docx"
    doc = Document()
    doc.add_heading(f"SOC nädala koondaruanne — {period_start} … {period_end}", level=1)

    # Lühike failiviide (vältimaks duplikaati TXT-ga)
    doc.add_heading("● Tekstiline kokkuvõte", level=2)
    doc.add_paragraph("(Allikafailide täielik loetelu ja perioodid — vt TXT aruannet.)")

    # Lisa TXT sisu dokumendi lõiku
    with open(out_txt, "r", encoding="utf-8") as fh:
        for line in fh:
            doc.add_paragraph(line.rstrip("\n"))

    doc.add_heading("● Graafikud ja visuaalid", level=2)
    images_week = [
        (trend_png, "Nädala trend: Alerts vs High+Critical"),
        (sev_png, "Severity jaotus päevade lõikes (stacked)"),
        (REP / f"week_action_bar_{today}.png", "Action jaotus (nädal, tulbad)"),
        (REP / f"week_action_pie_{today}.png", "Action osakaal (nädal, pirukas)"),
        (REP / f"week_threat_type_pie_{today}.png", "Threat/Content Type (nädal, pirukas)"),
        (REP / f"week_top_categories_{today}.png", "TOP 10 kategooriat (nädal)"),
        (REP / f"week_top_threats_{today}.png", "TOP 10 Threat / Content Name (nädal)"),
        (REP / f"week_top_src_{today}.png", "TOP 10 allika IP"),
        (REP / f"week_top_dst_{today}.png", "TOP 10 sihtmärgi IP"),
    ]
    for img, cap in images_week:
        add_image_safe(doc, img, cap, width_in=6.0)

    doc.save(str(docx_path))

    # --- Teated ---
    print("\n[OK] Nädala analüüs valmis.")
    print(f" - TXT : {out_txt}")
    print(f" - XLSX: {out_xlsx}")
    print(f" - DOCX: {docx_path}")
    print(f" - IP nimekiri: {ip_txt}")
    print(f" - Graafikud: {REP}")

if __name__ == "__main__":
    main()
