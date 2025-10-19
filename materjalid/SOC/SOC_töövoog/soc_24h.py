# ==============================================================
#  SOC 24h Logianalüüs v3.0.1
#  Autor: Heiki Rebane (õpiprojekt)
#  Kirjeldus:
#   - Loeb uusima CSV-faili kaustast raw/
#   - Koostab DOCX raporti graafikutega
#   - Lisab TOP 10 allika IP-d eraldi faili
#   - Märkib võimalikud valepositiivid
#   - Severity jaotus kuvatakse sõõrikuna legendiga
# ==============================================================

import re
from pathlib import Path
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

# --- Kaustad ---
BASE = Path.home() / "Documents" / "SOC"
RAW = BASE / "raw"
OUT = BASE / "tulemused"
REP = BASE / "reports"

for d in (RAW, OUT, REP):
    d.mkdir(parents=True, exist_ok=True)

# --- Lae threat-ID kirjeldused JSON-failist ---
with open("threat_descriptions.json", "r", encoding="utf-8") as f:
    THREAT_DESC = json.load(f)

# --- Abifunktsioonid ---
def iso_from_filename(name: str):
    """Ekstrakti kuupäev failinimest (dd.mm.yyyy või yyyy-mm-dd)"""
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", name)
    if m:
        return f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
    m2 = re.search(r"(\d{4})-(\d{2})-(\d{2})", name)
    return m2.group(0) if m2 else datetime.now().date().isoformat()

def first_existing(df, names):
    """Tagasta esimene veerunimi, mis eksisteerib andmestruktuuris"""
    for n in names:
        if n in df.columns:
            return n
    return None

def norm_lower(s):
    """Teisenda veerg väiketähtedeks ja eemalda tühikud"""
    return s.astype(str).str.strip().str.lower()

def bar(series, title, outpath, rot=0):
    """Loo tulpdiagramm ja salvesta PNG-failina"""
    if series is None or series.empty:
        return
    plt.figure(figsize=(10,5))
    series.plot(kind="bar", color="#888888")
    plt.title(title)
    plt.xticks(rotation=rot, ha="right" if rot else "center")
    plt.tight_layout()
    plt.savefig(outpath)
    plt.close()

def pie(series, title, outpath):
    """Loo sõõrikdiagramm ja salvesta PNG-failina"""
    if series is None or series.empty:
        return
    plt.figure(figsize=(7,7))
    wedges, texts, autotexts = plt.pie(
        series.values,
        labels=[str(i) for i in series.index],
        autopct="%1.1f%%",
        startangle=90,
        pctdistance=0.85,
        labeldistance=1.05
    )
    centre_circle = plt.Circle((0,0),0.70,fc='white')
    fig = plt.gcf()
    fig.gca().add_artist(centre_circle)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(outpath)
    plt.close()

def add_image(doc, img_path, caption="", width_in=6.0):
    """Lisa pilt DOCX-faili, kui see eksisteerib"""
    if img_path.exists():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_in))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

def mark_false_positive(row):
    """Märgista tõenäolised valepositiivid madal-risk kirjetel"""
    if row["sev_norm"] == "low":
        if row["Repeat Count"] > 200:
            if str(row["src_norm"]).startswith(("172.", "10.", "192.168.")):
                return True
        if "Nmap" in str(row["tname_norm"]) or "OpenSSL" in str(row["tname_norm"]):
            return True
    return False

def extract_threat_id(name):
    """Ekstrakti Threat-ID nimest"""
    match = re.search(r"\$\$(\d+)\$\$", str(name))
    return int(match.group(1)) if match else None

# --- Peafunktsioon ---
def main():
    # 1. Leia uusim CSV-fail
    csv_files = sorted(RAW.glob("*.csv"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not csv_files:
        print("[!] Ei leitud CSV-faile kaustas 'raw/'.")
        return

    path = csv_files[0]
    date_iso = iso_from_filename(path.name)
    print(f"[i] Analüüsitakse (uusim): {path.name} ({date_iso})")

    # 2. Loe CSV
    df = pd.read_csv(path, low_memory=False)

    # 3. Normaliseeri veerud
    sev_col = first_existing(df, ["Severity", "severity"])
    act_col = first_existing(df, ["Action", "action"])
    name_col = first_existing(df, ["Threat/Content Name", "Threat Name"])
    src_col = first_existing(df, ["Source address", "Source", "src"])
    cnt_col = first_existing(df, ["Repeat Count"])

    if sev_col: df["sev_norm"] = norm_lower(df[sev_col])
    if act_col: df["act_norm"] = norm_lower(df[act_col])
    if name_col: df["tname_norm"] = df[name_col].astype(str).str.strip()
    if src_col: df["src_norm"] = df[src_col].astype(str).str.strip()
    if cnt_col:
        df["Repeat Count"] = pd.to_numeric(df[cnt_col], errors="coerce").fillna(0).astype(int)

    # 4. Lisa threat-ID, kirjeldused ja valepositiivid
    df["threat_id"] = df["tname_norm"].apply(extract_threat_id)
    df["threat_desc"] = df["threat_id"].astype(str).map(THREAT_DESC).fillna("Kirjeldus puudub")
    df["false_positive"] = df.apply(mark_false_positive, axis=1)

    # 5. Arvuta statistika
    total = len(df)
    sev_counts = df["sev_norm"].value_counts()
    act_counts = df["act_norm"].value_counts()
    top_threats = df["tname_norm"].value_counts().head(10)
    top_src = df["src_norm"].value_counts().head(10)
    fp_count = df["false_positive"].sum()

    # 6. Graafikud
    pie(sev_counts, f"Severity jaotus – {date_iso}", REP / f"severity_donut_{date_iso}.png")
    bar(act_counts, f"Action jaotus – {date_iso}", REP / f"action_bar_{date_iso}.png")
    bar(top_threats, f"TOP 10 Threat – {date_iso}", REP / f"top_threats_{date_iso}.png", rot=45)
    bar(top_src, f"TOP allikad – {date_iso}", REP / f"top_src_{date_iso}.png", rot=45)

    # 7. DOCX raport
    doc = Document()
    doc.add_heading(f"SOC 24h Aruanne – {date_iso}", level=1)

    doc.add_heading("Ülevaade", level=2)
    doc.add_paragraph(f"- Logifail: {path.name}")
    doc.add_paragraph(f"- Kirjeid kokku: {total}")
    doc.add_paragraph(f"- Võimalikke valepositiive: {fp_count}")

    doc.add_heading("Graafikud", level=2)
    add_image(doc, REP / f"severity_donut_{date_iso}.png", "Severity jaotus (sõõrik)")
    add_image(doc, REP / f"action_bar_{date_iso}.png", "Action jaotus (tulp)")
    add_image(doc, REP / f"top_threats_{date_iso}.png", "TOP 10 Threat-ID")
    add_image(doc, REP / f"top_src_{date_iso}.png", "TOP allikad")

    doc.add_heading("TOP Threat-ID kirjeldused", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Nimi"
    hdr[2].text = "Kirjeldus"
    for name, count in top_threats.items():
        tid = extract_threat_id(name)
        desc = THREAT_DESC.get(str(tid), "Kirjeldus puudub")
        row = table.add_row().cells
        row[0].text = str(tid or "")
        row[1].text = name
        row[2].text = desc

    doc.add_heading("Valepositiivid", level=2)
    fp_df = df[df["false_positive"] == True]
    fp_top = fp_df["tname_norm"].value_counts().head(10)
    for name, count in fp_top.items():
        doc.add_paragraph(f"⚠️ {name} ({count} korda)")

    docx_path = OUT / f"24h_summary_{date_iso}.docx"
    doc.save(docx_path)

    # 8. TOP IP nimekiri
    ip_out = OUT / f"24h_ip_{date_iso}.txt"
    with open(ip_out, "w", encoding="utf-8") as f:
        f.write("TOP 10 Allika IP (24h)\n")
        f.write("="*30 + "\n\n")
        for ip, count in top_src.items():
            f.write(f"{ip} ({count} korda)\n")

    print(f"[OK] DOCX raport valmis: {docx_path}")
    print(f"[OK] IP nimekiri: {ip_out}")

if __name__ == "__main__":
    main()
