# ==============================================================
#  SOC 7-päevane Logianalüüs v3.0
#  Autor: Heiki Rebane (õpiprojekt)
#  Kirjeldus:
#   - Leiab ja kasutab 5–7 päevast logifailide andmestikku
#   - Koostab DOCX aruande ja graafikud
#   - TOP 10 Threat-ID ja IP-d
# ==============================================================

import re
from pathlib import Path
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

BASE = Path.home() / "Documents" / "SOC"
RAW = BASE / "raw"
OUT = BASE / "tulemused"
REP = BASE / "reports"

for d in (RAW, OUT, REP):
    d.mkdir(parents=True, exist_ok=True)

with open("threat_descriptions.json", "r", encoding="utf-8") as f:
    THREAT_DESC = json.load(f)

def iso_from_filename(name: str):
    """Ekstrakti kuupäev failinimest"""
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", name)
    if m:
        return f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
    m2 = re.search(r"(\d{4})-(\d{2})-(\d{2})", name)
    return m2.group(0) if m2 else datetime.now().date().isoformat()

def extract_threat_id(name):
    """Ekstrakti Threat-ID nimest"""
    match = re.search(r"\$\$(\d+)\$\$", str(name))
    return int(match.group(1)) if match else None

def main():
    # 1. Leia kuni 7 viimast faili
    csv_files = sorted(RAW.glob("*.csv"), key=lambda p: p.stat().st_mtime, reverse=True)[:7]
    if not csv_files:
        print("[!] Ei leitud CSV-faile kaustas 'raw/'.")
        return

    dfs = [pd.read_csv(p, low_memory=False) for p in csv_files]
    df = pd.concat(dfs, ignore_index=True)

    # 2. Otsi Threat veerg ja ekstrakti ID-d
    name_col = next((c for c in ["Threat/Content Name", "Threat Name"] if c in df.columns), None)
    if not name_col:
        print("[!] Threat veergu ei leitud.")
        return

    df["threat_id"] = df[name_col].apply(extract_threat_id)
    df["threat_desc"] = df["threat_id"].astype(str).map(THREAT_DESC).fillna("Kirjeldus puudub")
    top_threats = df[name_col].value_counts().head(10)

    # 3. Graafik ja DOCX
    plt.figure(figsize=(10,5))
    top_threats.plot(kind="bar", color="#888888")
    plt.title("TOP 10 Threat nädalas")
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    img_path = REP / "top_threats_week.png"
    plt.savefig(img_path)
    plt.close()

    doc = Document()
    doc.add_heading("SOC 7-päevane Aruanne", level=1)
    doc.add_paragraph(f"- Kasutatud {len(csv_files)} faili")
    doc.add_paragraph(f"- Kirjeid kokku: {len(df)}")

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx_path = OUT / "week_summary.docx"
    doc.save(docx_path)

    print(f"[OK] DOCX nädalakokkuvõte: {docx_path}")

if __name__ == "__main__":
    main()
