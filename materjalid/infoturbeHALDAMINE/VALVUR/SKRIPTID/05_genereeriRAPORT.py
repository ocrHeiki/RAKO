#!/usr/bin/env python3
"""
05_genereeri_raport.py - Koostab analüüsi tulemustest Wordi (.docx) raporti.
Teisendab UTC ajatemplid Eesti ajavööndisse (Europe/Tallinn).
Kasutamine: python3 SKRIPTID/05_genereeri_raport.py
"""

import os
import csv
from datetime import datetime
from zoneinfo import ZoneInfo
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ASCII Logo metainfo (VALVUR standard)
LOGO = r"""
###############################################################################
#                                                                             #
#   PROJEKT:     VALVUR - Raporti Genereerimine                               #
#   FAILI NIMI:  05_genereeriRAPORT.py                                       #
#   LOODUD:      27.03.2026                                                   #
#   AUTOR:       Heiki Rebane                                                 #
#   KIRJELDUS:   Koostab kronoloogilise raporti Eesti ajavööndis.             #
#                                                                             #
###############################################################################
"""

def convert_to_ee_time(utc_string):
    """
    Teisendab Windowsi UTC (Z) aja Eesti ajavööndisse.
    Näide: 2026-03-27T14:00:00.123Z -> 27.03.2026 16:00:00
    """
    if not utc_string or utc_string == "N/A":
        return "N/A"
    try:
        # Puhastame stringi ja muudame selle datetime objektiks
        # Windowsi ISO formaat võib sisaldada 'Z' lõpus
        clean_time = utc_string.replace('Z', '+00:00')
        dt_utc = datetime.fromisoformat(clean_time)
        
        # Teisendame Eesti aega (arvestab automaatselt suve/talveaega)
        dt_ee = dt_utc.astimezone(ZoneInfo("Europe/Tallinn"))
        
        return dt_ee.strftime('%d.%m.%Y %H:%M:%S')
    except Exception:
        return utc_string

def create_word_report():
    print(LOGO)
    doc = Document()

    # --- 1. TIITELLEHT JA PEALKIRI ---
    title = doc.add_heading('VALVUR - Intsidendi Analüüsi Raport', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    import socket
    hostname = socket.gethostname()
    run = meta.add_run(f"Auditeeritud masin: {hostname}\nRaporti koostamise aeg (Eesti): {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    run.italic = True

    # --- UUS: ISMS KIRJELDUS ---
    doc.add_heading('Infoturbe juhtimissüsteemi (ISMS) rakendamine', level=1)
    p = doc.add_paragraph("Näidisettevõttes rakendatakse infoturvet tuginedes E-ITS standardile. Juhtimissüsteem põhineb pideval parandamisel (PDCA tsükkel), kus teostatakse regulaarseid turvaauditeid (käesolev raport on osa sellest) ja riskihindamist. Tehnilised meetmed on kooskõlas riikliku infoturbe standardi baasturbe nõuetega.")

    # --- 2. KRONOLOOGILINE TIMELINE (Skriptist 02) ---
    doc.add_heading('1. Sündmuste kronoloogia (Timeline)', level=1)
    doc.add_paragraph("Kriitiliste sündmuste koondtabel Eesti ajavööndis (EET/EEST).")
    
    csv_path = 'TULEMUSED/02_tulemus_turvafiltreering.csv'
    if os.path.exists(csv_path):
        with open(csv_path, mode='r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Päis
            hdr_cells = table.rows[0].cells
            for i, text in enumerate(['Aeg (Eesti)', 'Event ID', 'Masin', 'Sündmuse sisu']):
                hdr_cells[i].text = text
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            for row in reader:
                row_cells = table.add_row().cells
                row_cells[0].text = convert_to_ee_time(row.get('TimeCreated', ''))
                row_cells[1].text = row.get('Id', 'N/A')
                row_cells[2].text = row.get('MachineName', 'N/A')
                # Võtame sõnumist esimesed 100 märki, et tabel ei läheks liiga laiaks
                msg = row.get('Message', 'N/A')
                row_cells[3].text = (msg[:100] + '...') if len(msg) > 100 else msg
    else:
        doc.add_paragraph("HOIATUS: Timeline faili ei leitud. Käivita skript 02.")

    # --- 3. TEOSTATUD OTSINGUD (Skriptist 03) ---
    doc.add_heading('2. Tuvastatud ründeindikaatorid (IOC)', level=1)
    keywords_path = 'TULEMUSED/03_tulemus_kahtlased_marksonad.csv'
    
    if os.path.exists(keywords_path):
        with open(keywords_path, mode='r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"MÄRKSÕNA: {row.get('MatchedKeyword')}").bold = True
                p.add_run(f" (MITRE: {row.get('MITRE_ID')}, CVE: {row.get('CVE_ID')})")
                p.add_run(f"\nKirjeldus: {row.get('Attack_Type', row.get('Description', ''))}")
                p.add_run(f"\nAeg: {convert_to_ee_time(row.get('TimeCreated'))}")
    else:
        doc.add_paragraph("Kahtlaseid märksõnu logidest ei leitud.")

    # --- 4. DEKODEERITUD POWERSHELL JA THREAT INTEL ---
    doc.add_heading('3. Süvaanalüüs ja Threat Intelligence', level=1)
    
    ti_path = 'TULEMUSED/09_tulemus_threat_intel.csv'
    if os.path.exists(ti_path):
        doc.add_paragraph("Tuvastatud välisühenduste maine kontroll:")
        with open(ti_path, mode='r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "IP Aadress"
            hdr[1].text = "Maine"
            hdr[2].text = "Soovitus"
            for row in reader:
                cells = table.add_row().cells
                cells[0].text = row.get('IP')
                cells[1].text = row.get('Reputation')
                cells[2].text = row.get('Action_Required')

    deep_path = 'TULEMUSED/04_tulemus_suvaanaluusi_raport.txt'
    
    if os.path.exists(deep_path):
        with open(deep_path, 'r', encoding='utf-8') as f:
            content = f.read()
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
    else:
        doc.add_paragraph("Süvaanalüüsi raportit ei leitud.")

    # --- 5. KAHTLASED FAILID JA ALLALAADIMISED (Skriptist 06) ---
    doc.add_heading('4. Kahtlased failid ja allalaadimised', level=1)
    files_path = 'TULEMUSED/06_tulemus_kahtlased_failid.csv'
    
    if os.path.exists(files_path):
        with open(files_path, mode='r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"PÕHJUS: {row.get('DetectionReason')}").bold = True
                p.add_run(f"\nAeg: {convert_to_ee_time(row.get('TimeCreated'))}")
                msg = row.get('Message', 'N/A')
                p.add_run(f"\nSisu: {(msg[:200] + '...') if len(msg) > 200 else msg}")
    else:
        doc.add_paragraph("Kahtlaseid failioperatsioone ei leitud.")

    # --- 6. E-ITS TURVAAUDIT (Skriptist 07) ---
    doc.add_heading('5. E-ITS Turvaaudit ja vastavuskontroll', level=1)
    audit_path = 'TULEMUSED/07_tulemus_turvaaudit.csv'
    
    if os.path.exists(audit_path):
        with open(audit_path, mode='r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Päis
            hdr_cells = table.rows[0].cells
            for i, text in enumerate(['Kontroll', 'Hetkeseis', 'Ootus', 'Staatus', 'Parandusmeede']):
                hdr_cells[i].text = text
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            for row in reader:
                row_cells = table.add_row().cells
                row_cells[0].text = row.get('Kontroll', '')
                row_cells[1].text = row.get('Hetkeseis', '')
                row_cells[2].text = row.get('Ootus', '')
                
                # Värvime FAIL staatuse punaseks
                status = row.get('Staatus', '')
                run = row_cells[3].paragraphs[0].add_run(status)
                if status == "FAIL":
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True
                
                row_cells[4].text = row.get('Meede', '')
    else:
        doc.add_paragraph("Auditi andmeid ei leitud.")

    # --- 7. VÕRGU SKANEERIMISE TULEMUSED (Skriptist 10) ---
    doc.add_heading('6. Tuvastatud võrguseadmed ja hostid', level=1)
    net_path = 'TULEMUSED/10_tulemus_vorgu_skaneerimine.txt'
    
    if os.path.exists(net_path):
        with open(net_path, 'r', encoding='utf-8') as f:
            content = f.read()
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
    else:
        doc.add_paragraph("Võrgu skaneerimise andmeid ei leitud.")

    # --- 8. TUVASTATUD KONTOD (Skriptist 11) ---
    doc.add_heading('7. Tuvastatud kasutajakontod', level=1)
    user_path = 'TULEMUSED/11_tulemus_kasutajad.txt'
    if os.path.exists(user_path):
        with open(user_path, 'r', encoding='utf-8') as f:
            doc.add_paragraph(f.read())

    # --- 9. INTSIDENDI TEAVITUSE MUSTAND ---
    doc.add_heading('8. Intsidendi teavituse mustand (Juhtkonnale)', level=1)
    notify = doc.add_paragraph()
    notify.add_run("TEEMA: KRIITILINE - Turvaintsidendi tuvastamine ja esmane reageerimine\n").bold = True
    notify.add_run("LUGUPEETUD JUHTKOND,\n\nTeavitame teid, et VALVUR seiresüsteem tuvastas kahtlase tegevuse meie sisevõrgus. ")
    notify.add_run("Esialgsel hinnangul on tegu ründega [SISESTA RÜNDE TÜÜP], mis puudutab [SISESTA MASINATE ARV] seadet. ")
    notify.add_run("Oleme rakendanud esmased piiramismeetmed ja jätkame süvaanalüüsiga. Täpsem detailne raport on lisatud käesoleva dokumendi tehnilistes sektsioonides.\n\n")
    notify.add_run("SOOVITUS: Rakendada paroolide sundvahetus ja kontrollida kriitiliste süsteemide varukoopiaid.").italic = True

    # --- SALVESTAMINE ---
    output_docx = 'TULEMUSED/VALVUR_LOPLIK_RAPORT.docx'
    try:
        doc.save(output_docx)
        print(f"\n[+] EDU: Raport salvestatud: {output_docx}")
    except Exception as e:
        print(f"\n[!] VIGA salvestamisel: {e}")

if __name__ == "__main__":
    create_word_report()
